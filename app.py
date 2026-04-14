import streamlit as st
import requests
import pandas as pd
import html
from io import BytesIO
import datetime
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from bs4 import BeautifulSoup
import calendar
import time
import re
from dateutil import parser
import urllib.parse
import cloudscraper
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ==========================================
# CONFIGURACIÓN INICIAL Y ESTILOS
# ==========================================
st.set_page_config(page_title="Boletín Mensual", layout="wide")

st.markdown("""
    <style>
    div.stButton > button, div.stDownloadButton > button {
        background-color: #00205B !important;
        color: white !important;
        border: none !important;
    }
    div.stButton > button:hover, div.stDownloadButton > button:hover {
        background-color: #00153D !important;
        color: white !important;
    }
    span[data-baseweb="tag"] {
        background-color: #00205B !important;
        color: white !important;
    }
    .github-footer {
        position: fixed;
        right: 20px;
        bottom: 20px;
        background-color: rgba(255, 255, 255, 0.9);
        padding: 8px 12px;
        border-radius: 50px;
        border: 1px solid #d0d7de;
        z-index: 1000;
        display: flex;
        align-items: center;
        font-family: 'Calibri', sans-serif;
        text-decoration: none;
        color: #24292f;
        box-shadow: 0px 4px 12px rgba(0,0,0,0.1);
        transition: transform 0.2s, box-shadow 0.2s;
    }
    .github-footer:hover {
        transform: translateY(-2px);
        box-shadow: 0px 6px 16px rgba(0,0,0,0.15);
        color: #00205B;
        border-color: #00205B;
    }
    .github-icon {
        margin-right: 8px;
        width: 22px;
        height: 22px;
    }
    </style>
    <a class="github-footer" href="https://github.com/sdiazprado" target="_blank">
        <img class="github-icon" src="https://github.githubassets.com/images/modules/logos_page/GitHub-Mark.png" alt="GitHub Logo">
        <span><strong>@sdiazprado</strong></span>
    </a>
""", unsafe_allow_html=True)

# ==========================================
# UTILIDADES DE FORMATO
# ==========================================
# ==========================================
# HERRAMIENTA DE RESCATE (TEXTO MANUAL)
# ==========================================
@st.cache_data(show_spinner=False)
def buscar_link_inteligente(titulo, organismo):
    """Cazador de DOIs de Doble Impacto (Estricto + Fuzzy). Cero Google."""
    import urllib.parse
    import requests
    import time
    import re

    # 1. Limpieza base
    titulo_raiz = re.split(r'[:\-]', titulo)[0].strip()
    titulo_limpio = re.sub(r'[^a-zA-Z0-9\s]', '', titulo_raiz)
    
    headers = {'User-Agent': 'mailto:bot_investigacion@banco.com'}
    time.sleep(0.5) 

    def consultar_api(query_param, texto_busqueda, modo_estricto=True):
        query_enc = urllib.parse.quote(texto_busqueda)
        url = f"https://api.crossref.org/works?{query_param}={query_enc}&select=URL,title,publisher&rows=4"
        
        try:
            res = requests.get(url, headers=headers, timeout=8)
            if res.status_code == 200:
                items = res.json().get('message', {}).get('items', [])
                
                for item in items:
                    url_oficial = item.get('URL')
                    if not url_oficial: continue
                        
                    pub = item.get('publisher', '').lower()
                    titulo_api = item.get('title', [''])[0].lower()
                    
                    if modo_estricto:
                        if 'oecd' in pub or 'organisation' in pub or organismo.lower() in pub:
                            return url_oficial
                    else:
                        titulo_comparar = titulo_limpio.lower()
                        if titulo_comparar in titulo_api or titulo_api in titulo_comparar:
                            return url_oficial
        except:
            pass
        return None

    link = consultar_api("query.title", titulo_limpio, modo_estricto=True)
    if link: return link

    time.sleep(0.5)
    link = consultar_api("query.bibliographic", titulo, modo_estricto=False)
    if link: return link

    return ""

def procesar_texto_pegado(texto_crudo, organismo_nombre):
    """Extrae Fecha y Título del texto pegado. Retorna DataFrame estandarizado."""
    rows = []
    lineas = [linea.strip() for linea in texto_crudo.split('\n') if linea.strip()]
    patron_fecha = r'(\d{1,2}\s+[A-Za-z]{3,}\s+\d{4})'
    
    i = 0
    while i < len(lineas):
        match_fecha = re.search(patron_fecha, lineas[i])
        if match_fecha:
            try:
                parsed_date = parser.parse(match_fecha.group(1))
            except:
                i += 1; continue
            
            titulo = ""
            if i >= 1:
                titulo = lineas[i-1]
                basura_menu = ['list view', 'grid view', 'z-a', 'a-z', 'oldest', 'most recent', 'most relevant', 'order by']
                if titulo.lower() in basura_menu and i >= 2: 
                    titulo = lineas[i-2]
            
            if titulo and len(titulo) > 10 and not any(b in titulo.lower() for b in ['search', 'filter', 'sort by', 'publications']):
                rows.append({
                    "Date": parsed_date, 
                    "Title": titulo,
                    "Link": "Pendiente",
                    "Organismo": organismo_nombre
                })
        i += 1
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.sort_values(by="Date", ascending=False).drop_duplicates(subset=['Title'])
    return df

def buscar_link_boe(titulo):
    """Busca silenciosamente en la web para obtener el Link Directo y Oficial del BoE"""
    import urllib.parse
    import requests
    from bs4 import BeautifulSoup
    import re
    
    # Extraemos solo el título limpio sin el autor para la búsqueda
    titulo_limpio = titulo.split(': ')[-1] if ': ' in titulo else titulo
    titulo_limpio = re.sub(r'[^a-zA-Z0-9\s]', '', titulo_limpio)
    
    # Usamos DuckDuckGo HTML para evadir bloqueos y obtener el link oficial sin usar Google
    query = f"site:bankofengland.co.uk/speech {titulo_limpio}"
    url = f"https://html.duckduckgo.com/html/?q={urllib.parse.quote(query)}"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
    
    try:
        res = requests.get(url, headers=headers, timeout=8)
        soup = BeautifulSoup(res.text, 'html.parser')
        
        # Atrapamos el link real de los resultados
        for a in soup.find_all('a', class_='result__url'):
            href = a.get('href', '').strip()
            if 'bankofengland.co.uk/speech' in href:
                if not href.startswith('http'):
                    href = 'https://' + href
                return href
    except:
        pass
        
    # Fallback de emergencia (1 clic)
    google_query = urllib.parse.quote(query)
    return f"https://www.google.com/search?q={google_query}"

def procesar_texto_pegado_boe(texto_crudo):
    """Extractor especializado para el formato del Bank of England (BoE)"""
    rows = []
    lineas = [linea.strip() for linea in texto_crudo.split('\n') if linea.strip()]
    patron_fecha = r'(\d{1,2}\s+[A-Za-z]{3,}\s+\d{4})'
    
    i = 0
    while i < len(lineas):
        match_fecha = re.search(patron_fecha, lineas[i])
        if match_fecha:
            try:
                parsed_date = parser.parse(match_fecha.group(1))
            except:
                i += 1; continue
            
            # 1. Buscar Autor un renglón ARRIBA (ej. "Speech // Phil Evans")
            autor = ""
            if i >= 1 and "//" in lineas[i-1]:
                partes = lineas[i-1].split("//")
                if len(partes) > 1:
                    autor = clean_author_name(partes[1].strip())
            
            # 2. Buscar Título Completo dos renglones ABAJO
            titulo = ""
            if i + 2 < len(lineas):
                titulo_raw = lineas[i+2]
                # Le quitamos el sufijo redundante " - speech by Autor"
                titulo_raw = re.sub(r'(?i)\s*[\-–—]\s*speech\s+by\s+.*$', '', titulo_raw).strip()
                titulo = titulo_raw
            
            # 3. Ensamblar y Guardar
            if titulo:
                titulo_final = f"{autor}: {titulo}" if autor else titulo
                rows.append({
                    "Date": parsed_date, 
                    "Title": titulo_final,
                    "Link": "Pendiente",
                    "Organismo": "BoE (Inglaterra)"
                })
        i += 1
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.sort_values(by="Date", ascending=False).drop_duplicates(subset=['Title'])
    return df
def clean_author_name(name):
    if not name: return ""
    minusc = ['de', 'van', 'von', 'der', 'del', 'la']
    words = name.strip().split()
    
    # Capitaliza todo excepto las preposiciones europeas
    cleaned_words = [w.capitalize() if w.lower() not in minusc else w.lower() for w in words]
    if cleaned_words:
        cleaned_words[0] = cleaned_words[0].capitalize() # La primera siempre mayúscula
        
    cleaned = " ".join(cleaned_words)
    # Arreglar iniciales pegadas (ej. "J.M. Keynes" -> "J. M. Keynes")
    cleaned = re.sub(r'\b([A-Z])\.\s*([A-Z])', lambda m: f"{m.group(1)}. {m.group(2)}", cleaned)
    return cleaned

# ==========================================
# FUNCIONES DE EXTRACCIÓN (BACKEND)
# ==========================================

# --- SECCIÓN: REPORTES ---
# BID (Annual Reports en inglés)
@st.cache_data(show_spinner=False)
def load_reportes_bid_en(start_date_str, end_date_str):
    """
    Extrae Annual Reports del BID en inglés usando cloudscraper
    (mismo método que funciona para BID Investigación)
    """
    import cloudscraper
    from bs4 import BeautifulSoup
    import datetime
    import re
    import time
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Reportes: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    page = 0
    max_pages = 5
    
    # Crear scraper con la misma configuración que usas en BID Investigación
    scraper = cloudscraper.create_scraper(
        browser={
            'browser': 'chrome',
            'platform': 'windows',
            'mobile': False
        },
        delay=5
    )
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.5',
        'Accept-Encoding': 'gzip, deflate, br',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1',
    }
    
    meses_map = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
        'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
        'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
    }
    
    while page < max_pages:
        url = f"https://publications.iadb.org/en?f%5B0%5D=type%3AAnnual%20Reports&page={page}"
        print(f"📄 Página {page+1}: {url}")
        
        try:
            response = scraper.get(url, headers=headers, timeout=30)
            
            if response.status_code != 200:
                print(f"   ❌ Error HTTP: {response.status_code}")
                break
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Buscar artículos
            items = soup.find_all('div', class_='views-row')
            
            if not items:
                print(f"   📭 No hay resultados en página {page+1}")
                # Guardar HTML para depuración
                with open(f"bid_reportes_page_{page}_debug.html", "w", encoding="utf-8") as f:
                    f.write(response.text)
                print(f"   💾 HTML guardado en bid_reportes_page_{page}_debug.html")
                break
            
            print(f"   📚 Artículos encontrados: {len(items)}")
            
            items_found = 0
            for item in items:
                try:
                    # Título y link
                    title_div = item.find('div', class_='views-field-field-title')
                    if not title_div:
                        continue
                    
                    a_tag = title_div.find('a')
                    if not a_tag:
                        continue
                    
                    titulo = a_tag.get_text(strip=True)
                    link = a_tag.get('href')
                    if link and not link.startswith('http'):
                        link = "https://publications.iadb.org" + link
                    
                    # Fecha
                    date_div = item.find('div', class_='views-field-field-date-issued-text')
                    if not date_div:
                        continue
                    
                    date_text = date_div.get_text(strip=True)
                    match = re.search(r'([A-Za-z]{3,9})\s+(\d{4})', date_text)
                    if not match:
                        continue
                    
                    mes_str = match.group(1).lower()[:3]
                    año = int(match.group(2))
                    mes_num = meses_map.get(mes_str, 1)
                    parsed_date = datetime.datetime(año, mes_num, 15)
                    
                    # Filtrar por fecha
                    if parsed_date < start_date or parsed_date > end_date:
                        continue
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "BID (Reportes)"
                        })
                        items_found += 1
                        print(f"   ✅ {parsed_date.date()} - {titulo[:50]}...")
                        
                except Exception as e:
                    continue
            
            print(f"   📊 Documentos en página {page+1}: {items_found}")
            
            if items_found == 0 and page > 0:
                break
            
            page += 1
            time.sleep(2)
            
        except Exception as e:
            print(f"   ❌ Error: {e}")
            break
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
    
    print(f"✅ BID Reportes - Total: {len(df)} documentos")
    return df

#Reportes - BPI - ESP -
@st.cache_data(show_spinner=False)
def load_reportes_bpi(start_date_str, end_date_str):
    """
    Extractor BPI - Reportes (BCBS, CPMI, IFC, CGFS)
    """
    import requests
    import pandas as pd
    import datetime
    import html
    from dateutil import parser
    from bs4 import BeautifulSoup
    import re
    import time
    
    # Configuración de fechas
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BPI Reportes: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    # ========== 1. URLs EXISTENTES (BCBS, CPMI) ==========
    urls_api = [
        "https://www.bis.org/api/document_lists/bcbspubls.json",
        "https://www.bis.org/api/document_lists/cpmi_publs.json"
    ]
    
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []

    # API calls (BCBS y CPMI) - TU CÓDIGO EXISTENTE
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            lista_documentos = data.get("list", {})
            for path, doc_info in lista_documentos.items():
                titulo = html.unescape(doc_info.get("short_title", ""))
                if not titulo:
                    continue
                link = "https://www.bis.org" + doc_info.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                date_str = doc_info.get("publication_start_date", "")
                parsed_date = None
                if date_str:
                    try:
                        parsed_date = parser.parse(date_str)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        pass
                if not parsed_date:
                    continue
                if start_date <= parsed_date <= end_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue

    # ========== 2. IFC publications (HTML) - TU CÓDIGO EXISTENTE ==========
    urls_html = ["https://www.bis.org/ifc/publications.htm"]
    for url in urls_html:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            content_div = soup.find('div', id='cmsContent')
            if not content_div:
                continue
            for p in content_div.find_all('p'):
                a_tag = p.find('a')
                if not a_tag:
                    continue
                titulo = a_tag.get_text(strip=True)
                href = a_tag.get('href', '')
                if not href or 'index.htm' in href:
                    continue
                link = "https://www.bis.org" + href if href.startswith('/') else href
                full_text = p.get_text(strip=True)
                date_str = full_text.replace(titulo, '').strip(', ')
                parsed_date = None
                if date_str:
                    try:
                        parsed_date = parser.parse(date_str)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        pass
                if not parsed_date:
                    match = re.search(r'\b(20\d{2})\b', titulo)
                    if match:
                        parsed_date = datetime.datetime(int(match.group(1)), 1, 1)
                if not parsed_date:
                    continue
                if start_date <= parsed_date <= end_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue

    # ========== 3. NUEVO: CGFS publications (API) ==========
    try:
        url_cgfs_api = "https://www.bis.org/api/document_lists/cgfs_publs.json"
        res = requests.get(url_cgfs_api, headers=headers, timeout=15)
        if res.status_code == 200:
            data = res.json()
            for path, doc in data.get("list", {}).items():
                titulo = html.unescape(doc.get("short_title", ""))
                if not titulo:
                    continue
                link = "https://www.bis.org" + doc.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                try:
                    parsed_date = parser.parse(doc.get("publication_start_date", ""))
                    if parsed_date.tzinfo is not None:
                        parsed_date = parsed_date.replace(tzinfo=None)
                    if start_date <= parsed_date <= end_date:
                        rows.append({"Date": parsed_date, "Title": titulo,
                                    "Link": link, "Organismo": "BPI"})
                except:
                    continue
    except Exception as e:
        print(f"   ⚠️ Error en CGFS API: {e}")

    # ========== 4. NUEVO: CGFS HTML (para documento No 71 que no está en API) ==========
    try:
        url_cgfs_html = "https://www.bis.org/cgfs_publs/index.htm"
        res = requests.get(url_cgfs_html, headers=headers, timeout=15)
        if res.status_code == 200:
            soup = BeautifulSoup(res.text, 'html.parser')
            
            # Buscar filas de la tabla
            rows_html = soup.find_all('tr')
            
            for row in rows_html:
                try:
                    cells = row.find_all('td')
                    if len(cells) < 2:
                        continue
                    
                    # Fecha en primera celda
                    date_text = cells[0].get_text(strip=True)
                    match = re.search(r'(\d{1,2})\s+([A-Za-z]{3,})\s+(\d{4})', date_text)
                    if not match:
                        continue
                    
                    day = int(match.group(1))
                    mes_str = match.group(2).lower()
                    año = int(match.group(3))
                    
                    meses = {'jan':1, 'feb':2, 'mar':3, 'apr':4, 'may':5, 'jun':6,
                            'jul':7, 'aug':8, 'sep':9, 'oct':10, 'nov':11, 'dec':12}
                    mes = meses.get(mes_str[:3], 1)
                    parsed_date = datetime.datetime(año, mes, day)
                    
                    if parsed_date < start_date or parsed_date > end_date:
                        continue
                    
                    # Título y enlace en segunda celda
                    link_elem = cells[1].find('a')
                    if not link_elem:
                        continue
                    
                    titulo = link_elem.get_text(strip=True)
                    link = link_elem.get('href')
                    if link and not link.startswith('http'):
                        link = "https://www.bis.org" + link
                    
                    # Evitar duplicados con los de la API
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "BPI"
                        })
                        print(f"   ✅ CGFS HTML: {parsed_date.date()} - {titulo[:50]}...")
                        
                except Exception as e:
                    continue
                    
    except Exception as e:
        print(f"   ⚠️ Error scraping CGFS: {e}")

    # ========== 5. Crear DataFrame final ==========
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None:
            df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    
    print(f"✅ BPI Reportes - Total final: {len(df)} documentos")
    return df


## Reportes BM
@st.cache_data(show_spinner=False)
def load_reportes_bm(start_date_str, end_date_str):
    """
    Extractor para Reportes del BM usando API de DSpace
    """
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}

    # ID exacto de la comunidad de Publicaciones
    scope_id = '06251f8a-62c2-59fb-add5-ec0993fc20d9'

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BM Reportes: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    # Palabras clave para identificar reportes (ampliadas)
    palabras_reporte = [
        r'\breport\b', r'\boutlook\b', r'\bprospects\b', r'\bupdate\b',
        r'\breview\b', r'\bmonitor\b', r'\bbulletin\b', r'\boverview\b',
        r'\bassessment\b', r'\banalysis\b', r'\bforecast\b', r'\btrends?\b',
        r'\bdevelopments?\b', r'\bglobal economic\b', r'\bcommodity markets\b',
        r'\beconomic\s+report\b', r'\bcountry\s+update\b', r'\bquarterly\b',
        r'\bannual\s+report\b', r'\bglobal\s+development\b', r'\bmacroeconomic\b',
        r'\bfiscal\s+update\b', r'\bpolicy\s+note\b', r'\bworking\s+paper\b',
        r'\bdiscussion\s+paper\b', r'\bpolicy\s+research\s+working\s+paper\b'
    ]

    rows = []
    page = 0
    max_pages = 10  # Aumentado para capturar más
    
    while page < max_pages:
        try:
            # Aumentar size a 50 para capturar más por página
            params = {
                'scope': scope_id,
                'sort': 'dc.date.issued,DESC',
                'page': page,
                'size': 50
            }
            res = requests.get(base_url, headers=headers, params=params, timeout=15)
            data = res.json()

            objects = data.get('_embedded', {}).get(
                'searchResult', {}).get('_embedded', {}).get('objects', [])
            
            if not objects:
                print(f"📭 No hay más resultados en página {page}")
                break

            print(f"📄 Página {page + 1}: {len(objects)} objetos encontrados")
            
            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})

                # Extraer Título
                title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                if not title:
                    continue
                
                # Extraer Fecha
                date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                if not date_s:
                    continue
                    
                try:
                    parsed_date = parser.parse(date_s)
                    if parsed_date.tzinfo is not None:
                        parsed_date = parsed_date.replace(tzinfo=None)
                except:
                    continue

                if parsed_date < start_date or parsed_date > end_date:
                    continue
                
                # Revisión de resultados 
                print(f"   📄 {parsed_date.date()} - {title[:80]}...")

                # ========== FILTRO MEJORADO ==========
                es_reporte = False
                
                # 1. Revisar título
                for palabra in palabras_reporte:
                    if re.search(palabra, title.lower()):
                        es_reporte = True
                        break
                
                # 2. Si no está en título, revisar descripción
                if not es_reporte:
                    abstract_list = meta.get('dc.description.abstract', [])
                    desc_list = meta.get('dc.description', [])
                    description = ""
                    if abstract_list:
                        description = abstract_list[0].get('value', '').lower()
                    elif desc_list:
                        description = desc_list[0].get('value', '').lower()
                    
                    for palabra in palabras_reporte:
                        if re.search(palabra, description):
                            es_reporte = True
                            break
                
                # 3. Si no es reporte, saltar
                #if not es_reporte:
                #    continue
                # ==================================(ESTE COMMENT evita que filtre innecesariamente todo el listado disponible)

                # Link permanente
                link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '')
                if not link:
                    link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"

                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date, 
                        "Title": title,
                        "Link": link, 
                        "Organismo": "BM"
                    })
                    items_found += 1
                    print(f"   ✅ {parsed_date.date()} - {title[:60]}...")

            print(f"   📊 Documentos en página {page + 1}: {items_found}")
            
            # Si no encontramos nada en 2 páginas consecutivas, paramos
            if items_found == 0 and page > 1:
                break
                
            page += 1
            time.sleep(0.5)
            
        except Exception as e:
            print(f"⚠️ Error en página {page}: {e}")
            break

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"✅ BM Reportes - Total: {len(df)} documentos")
    return df



## Reportes FM - 

@st.cache_data(show_spinner=False)
def load_reportes_fem(start_date_str, end_date_str):
    """Extractor FEM - Versión Selenium Final (Scroll + Fallback de Fecha)"""
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    import time
    import re

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://es.weforum.org/publications/"
    
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36")

    try:
        driver = webdriver.Chrome(options=chrome_options)
        driver.get(url)
        time.sleep(8)
        # Scroll para despertar la lista dinámica
        driver.execute_script("window.scrollTo(0, 1000);")
        time.sleep(4)
        
        js_script = """
        let res = [];
        document.querySelectorAll('a[href*="/publications/"]').forEach(el => {
            let title = el.innerText || el.textContent || "";
            let container = el.closest('article') || el.closest('div[class*="wef-"]') || el.parentElement;
            let date = container.querySelector('time')?.getAttribute('datetime');
            if (title.length > 15) {
                res.push({ t: title, l: el.href, d: date });
            }
        });
        return res;
        """
        extracted = driver.execute_script(js_script)
        driver.quit()

        for item in extracted:
            # Limpieza de título (quitar saltos de línea y frases de botones)
            titulo = item['t'].split('\n')[0]
            titulo = re.sub(r'(?i)Download PDF|Leer más|Read more|View details', '', titulo).strip()
            link = item['l']
            
            if "/series/" in link: continue

            # Parseo de Fecha
            parsed_date = None
            if item['d']:
                try: parsed_date = parser.parse(item['d']).replace(tzinfo=None)
                except: pass
            
            if not parsed_date:
                # Fallback: Extraer /YYYY/MM/ del link
                m = re.search(r'/(\d{4})/(\d{2})/', link)
                if m: parsed_date = datetime.datetime(int(m.group(1)), int(m.group(2)), 1)

            if parsed_date and start_date <= parsed_date <= end_date:
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "FEM"})
    except:
        pass

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False).drop_duplicates(subset=['Link'])
    return df

@st.cache_data(show_spinner=False)
def load_reportes_bm(start_date_str, end_date_str):
    """Extractor para Reportes del BM (Solo incluye los que mencionan 'Report')"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}
    
    # ID exacto de la comunidad compartida con Investigación
    scope_id = '06251f8a-62c2-59fb-add5-ec0993fc20d9'
    
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows, page = [], 0
    while True:
        try:
            params = {
                'scope': scope_id, 
                'sort': 'dc.date.issued,DESC', 
                'page': page, 
                'size': 20
            }
            res = requests.get(base_url, headers=headers, params=params, timeout=15)
            data = res.json()
            
            objects = data.get('_embedded', {}).get('searchResult', {}).get('_embedded', {}).get('objects', [])
            if not objects: break
            
            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})
                
                # Extraer Título y Fecha (Sin Autor, como acordamos)
                title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                
                parsed_date = None
                if date_s:
                    try: parsed_date = parser.parse(date_s)
                    except: pass
                
                if not parsed_date or parsed_date < start_date: continue
                
                # --- NUEVO FILTRO PRO-REPORTES ---
                abstract_list = meta.get('dc.description.abstract', [])
                desc_list = meta.get('dc.description', [])
                
                description = ""
                if abstract_list: description = abstract_list[0].get('value', '').lower()
                elif desc_list: description = desc_list[0].get('value', '').lower()
                
                # Si la palabra "report" NO está en la descripción, lo saltamos
                if not re.search(r'\breport\b', description):
                    continue
                # ----------------------------------
                
                # Link permanente
                link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '')
                if not link: link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"
                
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": title, "Link": link, "Organismo": "BM"})
                    items_found += 1
            
            if items_found == 0: break
            page += 1
            if page > 3: break # Límite para evitar búsquedas infinitas
            time.sleep(0.2)
        except:
            break
            
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_reportes_cef(start_date_str, end_date_str):
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        url = f"https://www.fsb.org/publications/?dps_paged={page}"
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            items = soup.find_all('div', class_=lambda c: c and 'post-excerpt' in c)
            if not items: break
            items_found = 0
            for item in items:
                title_div = item.find('div', class_='post-title')
                if not title_div or not title_div.find('a'): continue
                a_tag = title_div.find('a')
                titulo_raw = a_tag.get_text(strip=True)
                link = a_tag.get('href', '')
                date_div = item.find('div', class_='post-date')
                parsed_date = None
                if date_div:
                    try: parsed_date = parser.parse(date_div.get_text(strip=True))
                    except: pass
                if not parsed_date: continue
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "CEF"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.5) 
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

## Reportes OCDE 
@st.cache_data(show_spinner=False)
def load_reportes_ocde(start_date_str, end_date_str):
    """Extractor OCDE - Reports (API oficial)"""
    import requests
    import datetime
    import re
    import time
    from dateutil import parser

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 OCDE Reportes: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []

    # API base de la OCDE
    base_url = "https://api.oecd.org/webcms/search/faceted-search"

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json"
    }

    page = 0
    page_size = 50  # Número de resultados por página
    max_pages = 10  # Límite de seguridad
    documentos_procesados = 0

    print("📡 Solicitando Reportes a la API de la OCDE (con paginación)...")

    try:
        while page < max_pages:
            # Parámetros para buscar Reports en inglés
            params = {
                "siteName": "oecd",
                "interfaceLanguage": "en",
                "orderBy": "mostRecent",
                "pageSize": page_size,
                "page": page,
                "facets": "oecd-languages:en",
                "hiddenFacets": "oecd-content-types:publications/reports"  # <-- FILTRO PARA REPORTES
            }

            print(f"   📄 Procesando página {page + 1}...")
            response = requests.get(base_url, params=params, headers=headers, timeout=15)

            if response.status_code != 200:
                print(f"   ❌ Error en página {page + 1}: {response.status_code}")
                break

            data = response.json()

            # Buscar los resultados
            results = data.get("results", [])

            if not results:
                print(f"   📭 No hay más resultados en página {page + 1}")
                break

            documentos_en_pagina = 0
            fecha_mas_antigua = None

            for item in results:
                titulo = item.get("title", "") or item.get("name", "")
                link = item.get("url", "") or item.get("link", "")

                if not titulo or not link:
                    continue

                # Extraer fecha
                fecha_texto = item.get("publicationDateTime", "")
                parsed_date = None
                if fecha_texto:
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        continue

                if not parsed_date:
                    continue

                fecha_mas_antigua = parsed_date

                # Si el documento es más antiguo que start_date, paramos
                if parsed_date < start_date:
                    print(f"   ⏹️ Documento más antiguo que {start_date.strftime('%Y-%m')}, deteniendo paginación")
                    page = max_pages
                    break

                # Filtrar por rango de fechas
                if parsed_date >= start_date and parsed_date <= end_date:
                    # Limpiar título
                    titulo = re.sub(r'\s+', ' ', titulo).strip()

                    # Asegurar URL absoluta
                    if link.startswith('/'):
                        link = f"https://www.oecd.org{link}"

                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "OCDE"
                    })
                    documentos_en_pagina += 1
                    documentos_procesados += 1

            print(f"   📊 Página {page + 1}: {documentos_en_pagina} documentos en el rango")

            # Si no encontramos documentos en esta página y ya pasamos la fecha límite
            if documentos_en_pagina == 0 and fecha_mas_antigua and fecha_mas_antigua < start_date:
                print(f"   ⏹️ Fin de resultados para el mes solicitado")
                break

            # Si encontramos menos de page_size documentos, probablemente es la última página
            if len(results) < page_size:
                print(f"   📭 Última página alcanzada")
                break

            page += 1
            time.sleep(0.3)  # Pequeña pausa para no sobrecargar la API

        print(f"\n📊 Total Reportes OCDE encontrados: {documentos_procesados}")

    except Exception as e:
        print(f"❌ Error en load_reportes_ocde: {e}")
        import traceback
        traceback.print_exc()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])

    print(f"📊 OCDE Reportes - Total final: {len(df)}")
    return df
    



# --- SECCIÓN: PUBLICACIONES INSTITUCIONALES ---
@st.cache_data(show_spinner=False)
def load_pub_inst_ocde(start_date_str, end_date_str):
    """Extractor OCDE - Publicaciones Institucionales (API oficial)"""
    import requests
    import datetime
    import re
    import time
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 OCDE Pub. Institucionales: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    # API base de la OCDE
    base_url = "https://api.oecd.org/webcms/search/faceted-search"
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json"
    }
    
    page = 0
    page_size = 50
    max_pages = 10
    
    print("📡 Solicitando Publicaciones Institucionales a la API de la OCDE (con paginación)...")
    
    try:
        while page < max_pages:
            # Parámetros para buscar el sub-tema psi114
            params = {
                "siteName": "oecd",
                "interfaceLanguage": "en",
                "orderBy": "mostRecent",
                "pageSize": page_size,
                "page": page,
                "facets": "oecd-languages:en",
                "hiddenFacets": "oecd-policy-subissues:psi114"  # <-- FILTRO PARA PUB. INSTITUCIONALES
            }
            
            print(f"   📄 Procesando página {page + 1}...")
            response = requests.get(base_url, params=params, headers=headers, timeout=15)
            
            if response.status_code != 200:
                print(f"   ❌ Error en página {page + 1}: {response.status_code}")
                break
            
            data = response.json()
            
            # Buscar los resultados
            results = data.get("results", [])
            
            if not results:
                print(f"   📭 No hay más resultados en página {page + 1}")
                break
            
            documentos_en_pagina = 0
            fecha_mas_antigua = None
            
            for item in results:
                titulo = item.get("title", "") or item.get("name", "")
                link = item.get("url", "") or item.get("link", "")
                
                if not titulo or not link:
                    continue
                
                # Extraer fecha
                fecha_texto = item.get("publicationDateTime", "")
                
                parsed_date = None
                if fecha_texto:
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        continue
                
                if not parsed_date:
                    continue
                
                fecha_mas_antigua = parsed_date
                
                # Si el documento es más antiguo que start_date, paramos
                if parsed_date < start_date:
                    print(f"   ⏹️ Documento más antiguo que {start_date.strftime('%Y-%m')}, deteniendo paginación")
                    page = max_pages
                    break
                
                # Filtrar por rango de fechas
                if parsed_date >= start_date and parsed_date <= end_date:
                    # Limpiar título
                    titulo = re.sub(r'\s+', ' ', titulo).strip()
                    
                    # Asegurar URL absoluta
                    if link.startswith('/'):
                        link = f"https://www.oecd.org{link}"
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "OCDE"
                    })
                    documentos_en_pagina += 1
            
            print(f"   📊 Página {page + 1}: {documentos_en_pagina} documentos en el rango")
            
            # Si no encontramos documentos en esta página y ya pasamos la fecha límite
            if documentos_en_pagina == 0 and fecha_mas_antigua and fecha_mas_antigua < start_date:
                print(f"   ⏹️ Fin de resultados para el mes solicitado")
                break
            
            # Si encontramos menos de page_size documentos, probablemente es la última página
            if len(results) < page_size:
                print(f"   📭 Última página alcanzada")
                break
            
            page += 1
            time.sleep(0.3)
        
        print(f"\n📊 Total documentos OCDE Pub. Institucionales encontrados: {len(rows)}")
        
    except Exception as e:
        print(f"❌ Error en load_pub_inst_ocde: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 OCDE Pub. Institucionales - Total final: {len(df)}")
    return df

# --- Publicaciones Institucionales --- OEI 
@st.cache_data(show_spinner=False)
def load_pub_inst_oei(start_date_str, end_date_str):
    """Extractor OEI (IEO-IMF) - Versión API Next.js con headers completos"""
    import requests
    import datetime
    import re
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 OEI: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    # Intentar obtener el build ID dinámicamente desde la página HTML
    build_id = "qchYZivFKVMGvRneSTtnM"  # Fallback
    try:
        headers_browser = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
            'Sec-Fetch-Dest': 'document',
            'Sec-Fetch-Mode': 'navigate',
            'Sec-Fetch-Site': 'none',
            'Sec-Fetch-User': '?1',
            'Cache-Control': 'max-age=0',
        }
        res = requests.get("https://ieo.imf.org/en/Publications/annual-reports", headers=headers_browser, timeout=15)
        # Buscar el build ID en el HTML
        match = re.search(r'/_next/data/([a-zA-Z0-9]+)/en/publications/annual-reports\.json', res.text)
        if match:
            build_id = match.group(1)
            print(f"🔧 Build ID encontrado: {build_id}")
    except Exception as e:
        print(f"⚠️ Usando build ID por defecto: {build_id}")
    
    # URL del JSON
    url = f"https://ieo.imf.org/_next/data/{build_id}/en/publications/annual-reports.json"
    
    # Headers completos para simular un navegador real
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'application/json, text/plain, */*',
        'Accept-Language': 'en-US,en;q=0.5',
        'Accept-Encoding': 'gzip, deflate, br',
        'Referer': 'https://ieo.imf.org/en/Publications/annual-reports',
        'Origin': 'https://ieo.imf.org',
        'Connection': 'keep-alive',
        'Sec-Fetch-Dest': 'empty',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Site': 'same-origin',
        'Cache-Control': 'no-cache',
        'Pragma': 'no-cache',
    }
    
    try:
        print(f"📡 Consultando: {url}")
        response = requests.get(url, headers=headers, timeout=15)
        
        if response.status_code == 200:
            data = response.json()
            
            # ✅ RUTA CORRECTA según el JSON que analizamos
            try:
                # Los reportes están en componentProps.[id].fields.datasource.reports.results
                component_props = data.get('pageProps', {}).get('componentProps', {})
                
                # Buscar el componente ReportsListing
                reports_results = None
                for comp_id, comp_value in component_props.items():
                    if 'fields' in comp_value and 'datasource' in comp_value['fields']:
                        datasource = comp_value['fields']['datasource']
                        if 'reports' in datasource and 'results' in datasource['reports']:
                            reports_results = datasource['reports']['results']
                            print(f"✅ Componente encontrado: {comp_id}")
                            break
                
                if not reports_results:
                    print("⚠️ No se encontraron reportes en componentProps")
                    return pd.DataFrame()
                
                print(f"📚 Reportes encontrados: {len(reports_results)}")
                
                for report in reports_results:
                    # Extraer título
                    titulo = report.get('title', {}).get('jsonValue', {}).get('value', '')
                    
                    # Extraer fecha
                    fecha_texto = report.get('publicationDate', {}).get('jsonValue', {}).get('value', '')
                    
                    # Extraer link del PDF
                    completed_link = report.get('completedReportLink', {}).get('jsonValue', {}).get('value', {})
                    link = completed_link.get('href', '') if isinstance(completed_link, dict) else ''
                    
                    if not titulo or not fecha_texto:
                        continue
                    
                    # Parsear fecha
                    parsed_date = parser.parse(fecha_texto).replace(tzinfo=None)
                    
                    if parsed_date < start_date or parsed_date > end_date:
                        continue
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "OEI"
                    })
                    print(f"   ✅ {parsed_date.strftime('%Y-%m-%d')}: {titulo}")
                
            except Exception as e:
                print(f"   ⚠️ Error procesando el JSON: {e}")
                import traceback
                traceback.print_exc()
        else:
            print(f"❌ Error en la API: {response.status_code}")
            print(f"   Respuesta: {response.text[:200] if response.text else 'Vacía'}")
            
    except Exception as e:
        print(f"❌ Error en load_pub_inst_oei: {e}")
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 OEI - Total documentos: {len(df)}")
    return df

# ========== FUNCIÓN PARA CEMLA (PUBLICACIONES INSTITUCIONALES) ==========
@st.cache_data(show_spinner=False)
def load_pub_inst_cemla(start_date_str, end_date_str):
    """
    Extractor CEMLA - Publicaciones Institucionales (Novedades individuales)
    Filtra eventos y contenido no académico
    """
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import re
    import pandas as pd
    import time
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    print("="*50)
    print("🔍 CEMLA PUBLICACIONES - Extrayendo novedades de boletines...")
    print(f"📅 Rango solicitado: {start_date_str} a {end_date_str}")
    print("="*50)

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"✅ Fechas parseadas: {start_date.date()} a {end_date.date()}")
    except Exception as e:
        print(f"⚠️ Error parseando fechas: {e}")
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now() + datetime.timedelta(days=365)

    # Palabras a excluir (eventos, cursos, etc. - no publicaciones académicas)
    palabras_excluir = [
        'reunión', 'reunion', 'virtual', 'curso', 'taller', 'seminario',
        'conferencia', 'webinar', 'congreso', 'foro', 'encuentro',
        'junta', 'comité', 'comite', 'próximas actividades', 'calendario',
        'convocatoria', 'premio', 'inscripción', 'registro'
    ]

    rows = []
    
    url = "https://www.cemla.org/comunicados.html"
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
    }

    try:
        # =========================================================
        # PASO 1: Obtener la lista de boletines
        # =========================================================
        print(f"📡 Solicitando página de boletines: {url}")
        response = requests.get(url, headers=headers, timeout=30, verify=False)
        print(f"   Status Code: {response.status_code}")
        
        if response.status_code != 200:
            print(f"❌ Error al acceder a la página")
            return pd.DataFrame()

        soup = BeautifulSoup(response.text, 'html.parser')
        
        meses = {
            'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4,
            'mayo': 5, 'junio': 6, 'julio': 7, 'agosto': 8,
            'septiembre': 9, 'octubre': 10, 'noviembre': 11, 'diciembre': 12
        }
        
        # Encontrar todos los boletines en el rango de fechas
        boletines_a_procesar = []
        
        for ul in soup.find_all('ul', class_='iconlist'):
            for li in ul.find_all('li'):
                p = li.find('p')
                if not p:
                    continue
                
                a_tag = p.find('a')
                if not a_tag:
                    continue
                
                titulo_texto = a_tag.get_text(strip=True)
                link = a_tag.get('href', '')
                
                match = re.match(r'^([A-Za-z]+)\s+(\d{4})$', titulo_texto, re.IGNORECASE)
                if match:
                    mes_str, año = match.groups()
                    mes_num = meses.get(mes_str.lower(), 0)
                    
                    if mes_num:
                        fecha = datetime.datetime(int(año), mes_num, 1)
                        
                        if start_date <= fecha <= end_date:
                            boletines_a_procesar.append({
                                'fecha': fecha,
                                'titulo': titulo_texto,
                                'link': link
                            })
                            print(f"📌 Boletín encontrado: {fecha.strftime('%Y-%m')} - {titulo_texto}")
        
        print(f"✅ Total boletines en rango: {len(boletines_a_procesar)}")
        
        if not boletines_a_procesar:
            print("⚠️ No se encontraron boletines en el rango de fechas")
            return pd.DataFrame()
        
        # =========================================================
        # PASO 2: Procesar cada boletín y extraer sus novedades
        # =========================================================
        for boletin in boletines_a_procesar:
            print(f"\n🔍 Procesando boletín: {boletin['titulo']} ({boletin['link']})")
            
            try:
                time.sleep(1)
                
                res_boletin = requests.get(boletin['link'], headers=headers, timeout=30, verify=False)
                if res_boletin.status_code != 200:
                    print(f"  ⚠️ Error al acceder al boletín: {res_boletin.status_code}")
                    continue
                
                soup_boletin = BeautifulSoup(res_boletin.text, 'html.parser')
                
                # Buscar todas las novedades (divs con clase "ipost clearfix")
                novedades = soup_boletin.find_all('div', class_=lambda c: c and 'ipost' in c.split() if c else False)
                
                if not novedades:
                    print(f"  ⚠️ No se encontraron novedades en este boletín")
                    continue
                
                print(f"  📚 Novedades encontradas: {len(novedades)}")
                
                for novedad in novedades:
                    try:
                        # Extraer título
                        title_elem = novedad.find('div', class_='entry-title')
                        if not title_elem:
                            continue
                        
                        h3 = title_elem.find('h3')
                        if not h3:
                            continue
                        
                        titulo = h3.get_text(strip=True)
                        
                        # ===== FILTRO: Excluir eventos y contenido no académico =====
                        titulo_lower = titulo.lower()
                        es_excluido = any(palabra in titulo_lower for palabra in palabras_excluir)
                        
                        if es_excluido:
                            print(f"    ⏭️ Excluido (evento): {titulo[:60]}...")
                            continue
                        
                        # Extraer descripción y enlace
                        content_elem = novedad.find('div', class_='entry-content')
                        if not content_elem:
                            continue
                        
                        p = content_elem.find('p')
                        if not p:
                            continue
                        
                        # Extraer el enlace "Leer más..."
                        a_link = p.find('a', href=True)
                        if a_link:
                            link_novedad = a_link.get('href', '')
                            descripcion = p.get_text(strip=True).replace(a_link.get_text(strip=True), '').strip()
                        else:
                            link_novedad = boletin['link']
                            descripcion = p.get_text(strip=True)
                        
                        # Limpiar título
                        titulo = re.sub(r'\s+', ' ', titulo).strip()
                        
                        # Solo agregar si el título es significativo
                        if titulo and len(titulo) > 10:
                            rows.append({
                                'Date': boletin['fecha'],
                                'Title': titulo,
                                'Link': link_novedad if link_novedad else boletin['link'],
                                'Organismo': "CEMLA"
                            })
                            print(f"    ✅ {titulo[:60]}...")
                    
                    except Exception as e:
                        print(f"    ⚠️ Error procesando novedad: {e}")
                        continue
                        
            except Exception as e:
                print(f"  ❌ Error procesando boletín: {e}")
                continue

    except Exception as e:
        print(f"❌ Error general: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'], keep='first')
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ CEMLA PUBLICACIONES - Total novedades: {len(df)} documentos")

    return df

# -- G20 --
@st.cache_data(show_spinner=False)
def load_pub_inst_g20(start_date_str, end_date_str):
    """Extrae documentos del G20 desde la página de News and Media"""
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import re
    import pandas as pd

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 G20: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")

    url = "https://g20.org/media/"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    
    # Palabras clave que queremos incluir
    keywords_incluir = [
        'chair summary', 'declarations', 'g-20 note', 'presidency note',
        'chair\'s summary', 'chair summary', 'g20 note', 'presidency note'
    ]
    
    # Palabras clave para excluir
    keywords_excluir = [
        'agriculture', 'cultura', 'cultural', 'food security', 
        'farming', 'rural', 'agri'
    ]
    
    rows = []

    try:
        print(f"📡 Solicitando página: {url}")
        res = requests.get(url, headers=headers, timeout=15)
        
        if res.status_code != 200:
            print(f"❌ Error al acceder a la página: {res.status_code}")
            return pd.DataFrame()

        soup = BeautifulSoup(res.text, 'html.parser')
        
        # Buscar la sección de Press Releases
        press_section = None
        for section in soup.find_all('section', class_='paragraphsection'):
            toptitle = section.find('h2', class_='toptitle')
            if toptitle and 'Press Releases' in toptitle.get_text():
                press_section = section
                break
        
        if not press_section:
            print("⚠️ No se encontró la sección de Press Releases")
            return pd.DataFrame()
        
        # Buscar todos los artículos (h2 seguido de p con fecha)
        articles = press_section.find_all(['h2', 'p'])
        
        i = 0
        while i < len(articles) - 1:
            if articles[i].name == 'h2':
                h2 = articles[i]
                a_tag = h2.find('a')
                
                if a_tag and a_tag.get('href'):
                    titulo = a_tag.get_text(strip=True)
                    link = a_tag.get('href')
                    
                    if not titulo:
                        i += 1
                        continue
                    
                    if i + 1 < len(articles) and articles[i + 1].name == 'p':
                        p_text = articles[i + 1].get_text(strip=True)
                        
                        match = re.search(r'([A-Za-z]+ \d{1,2},? \d{4})', p_text)
                        if match:
                            fecha_str = match.group(1)
                            try:
                                fecha_str = fecha_str.replace(',', '')
                                parsed_date = datetime.datetime.strptime(fecha_str, '%B %d %Y')
                            except:
                                try:
                                    parsed_date = datetime.datetime.strptime(fecha_str, '%b %d %Y')
                                except:
                                    parsed_date = None
                            
                            if parsed_date:
                                if parsed_date < start_date or parsed_date > end_date:
                                    i += 2
                                    continue
                                
                                titulo_lower = titulo.lower()
                                incluir = any(kw in titulo_lower for kw in keywords_incluir)
                                excluir = any(kw in titulo_lower for kw in keywords_excluir)
                                
                                if excluir or not incluir:
                                    i += 2
                                    continue
                                
                                if link.startswith('/'):
                                    link = f"https://g20.org{link}"
                                
                                rows.append({
                                    "Date": parsed_date,
                                    "Title": titulo,
                                    "Link": link,
                                    "Organismo": "G20"
                                })
                                print(f"   ✅ Agregado: {titulo[:60]}... ({parsed_date.date()})")
                    else:
                        print(f"   ⚠️ No hay párrafo después del h2")
                else:
                    print(f"   ⚠️ h2 sin enlace válido")
            
            i += 1
            
    except Exception as e:
        print(f"❌ Error general: {e}")
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ TOTAL G20: {len(df)} documentos")
    else:
        print("⚠️ No se encontraron documentos del G20")

    return df

@st.cache_data(show_spinner=False)
def load_pub_inst_cef(start_date_str, end_date_str):
    url = "https://www.fsb.org/publications/key-regular-publications/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows = []
    try:
        res = requests.get(url, headers=headers, timeout=15)
        soup = BeautifulSoup(res.text, 'html.parser')
        for section in soup.find_all('div', class_='wp-bootstrap-blocks-row'):
            h2 = section.find('h2')
            if not h2: continue
            base_title = h2.get_text(strip=True)
            # Latest
            latest_btn = section.find('button', class_='btn-primary')
            if latest_btn and latest_btn.find('a'):
                a_tag = latest_btn.find('a')
                link = "https://www.fsb.org" + a_tag['href'] if a_tag['href'].startswith('/') else a_tag['href']
                date_match = re.search(r'\((.*?)\)', a_tag.get_text())
                parsed_date = parser.parse(date_match.group(1)) if date_match else None
                if parsed_date and parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": f"{base_title}: Latest Report", "Link": link, "Organismo": "CEF"})
            # Previous
            dropdown = section.find('div', class_='dropdown-menu')
            if dropdown:
                for l in dropdown.find_all('a'):
                    year_text = l.get_text(strip=True)
                    try: parsed_date = datetime.datetime(int(year_text), 1, 1)
                    except: parsed_date = None
                    if parsed_date and parsed_date >= start_date:
                        rows.append({"Date": parsed_date, "Title": f"{base_title} ({year_text})", "Link": l['href'], "Organismo": "CEF"})
    except: pass
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# ========== FUNCIÓN UNIVERSAL PARA NOTICIAS DEL FMI (API COVEO) ==========
@st.cache_data(show_spinner=False)
def load_fmi_news_all(start_date_str, end_date_str):
    """
    Extrae TODAS las noticias del FMI usando la API de Coveo.
    Incluye Press Releases, Mission Concluding, Statements, News Articles, etc.
    """
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 FMI News (API Coveo): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")

    rows = []
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"

    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",
        "Referer": "https://www.imf.org/",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    # Filtro para capturar TODO excepto discursos
    payload = {
        "aq": "@imftype==(\"News Article\",\"Press Release\",\"Communique\",\"Mission Concluding Statement\",\"News Brief\",\"Public Information Notice\",\"Statements at Donor Meeting\",\"Views and Commentaries\",\"Blog Page\",\"IMF Staff Country Reports\") AND NOT @imftype==(\"Speech\",\"Transcript\") AND @syslanguage==\"English\"",
        "numberOfResults": 300,
        "sortCriteria": "@imfdate descending"
    }

    try:
        print("📡 Solicitando noticias del FMI a la API de Coveo...")
        # Deshabilitar verificación SSL para evitar errores locales
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        res = requests.post(url, headers=headers, json=payload, timeout=15, verify=False)

        if res.status_code == 200:
            data = res.json()
            print(f"✅ Respuesta recibida. Total en API: {data.get('totalCount', 0)} resultados")

            for item in data.get("results", []):
                titulo = item.get("title", "").strip()
                link = item.get("clickUri", "")
                content_type = item.get("raw", {}).get("imftype", "Unknown")

                raw_date = item.get("raw", {}).get("date")
                parsed_date = None
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except:
                        pass

                if not titulo or not link or not parsed_date:
                    continue

                if start_date <= parsed_date <= end_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "FMI"
                        })
                        print(f"   ✅ [{content_type[:25]}] {parsed_date.strftime('%Y-%m-%d')}: {titulo[:60]}...")
        else:
            print(f"❌ Error en la API: {res.status_code}")
            print(f"   Respuesta: {res.text[:200]}")

    except Exception as e:
        print(f"❌ Error: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ TOTAL FMI News (API): {len(df)} documentos")
    else:
        print("⚠️ No se encontraron documentos")

    return df

# -- BPI -- Publicaciones Institucionales 
@st.cache_data(show_spinner=False)
def load_pub_inst_bpi(start_date_str, end_date_str):
    urls_api = ["https://www.bis.org/api/document_lists/annualeconomicreports.json", "https://www.bis.org/api/document_lists/quarterlyreviews.json"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows = []
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            for path, doc in data.get("list", {}).items():
                titulo = html.unescape(doc.get("short_title", ""))
                link = "https://www.bis.org" + doc.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"): link += ".htm"
                try: parsed_date = parser.parse(doc.get("publication_start_date", ""))
                except: continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "BPI"})
        except: continue
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_country_reports_fmi(start_date_str, end_date_str):
    """Extractor FMI - Country Reports (Conexión Directa a Coveo API)"""
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    # 1. EL ENDPOINT Y LA LLAVE MAESTRA QUE DESCUBRISTE
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }
    
    # 2. EL PAYLOAD (Falsificamos la petición del buscador)
    payload = {
        "aq": "@imfseries==\"IMF Staff Country Reports\"", # Filtro estricto por la Serie
        "numberOfResults": 100, # Cantidad a traer (Suficiente para un mes)
        "sortCriteria": "@imfdate descending" # Los más recientes primero
    }
    
    try:
        # Hacemos un POST directo a la base de datos de Coveo
        res = requests.post(url, headers=headers, json=payload, timeout=15)
        
        if res.status_code == 200:
            data = res.json()
            
            # 3. EXTRACCIÓN (Limpia y sin HTML)
            for item in data.get("results", []):
                titulo = item.get("title", "")
                link = item.get("clickUri", "")
                
                # La fecha viene en timestamp (milisegundos). Lo dividimos entre 1000 para segundos.
                raw_date = item.get("raw", {}).get("date")
                parsed_date = None
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except: pass
                
                if not titulo or not link or not parsed_date: continue
                
                # Validamos contra la fecha del filtro de la app
                if parsed_date >= start_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "FMI"})
    except Exception as e:
        pass
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_press_releases_fmi(start_date_str, end_date_str):
    """Extractor FMI - Press Releases (Historial completo vía Coveo API)"""
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    # 1. El Endpoint y la llave que tú mismo descubriste
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    
    # 2. Inyección de Headers para evadir el bloqueo CORS
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",   # <--- LA LLAVE PARA ENTRAR
        "Referer": "https://www.imf.org/", # <--- CONFIRMA QUE "VENIMOS" DEL FMI
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }
    
    # 3. Payload: Agregamos el filtro estricto de idioma
    payload = {
        # Le pedimos PRs Y que el idioma sea inglés
        "aq": "@imftype==\"Press Release\" AND @syslanguage==\"English\"", 
        "numberOfResults": 150, 
        "sortCriteria": "@imfdate descending"
    }
    
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=15)
        
        if res.status_code == 200:
            data = res.json()
            
            for item in data.get("results", []):
                titulo = item.get("title", "")
                link = item.get("clickUri", "")
                
                # Coveo entrega la fecha en formato Unix (Milisegundos). 
                # ¡Es perfecto porque no falla la conversión!
                raw_date = item.get("raw", {}).get("date")
                parsed_date = None
                if raw_date:
                    try:
                        # Convertimos de milisegundos a fecha normal
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except: pass
                
                if not titulo or not link or not parsed_date: continue
                
                # Filtro final de fechas
                if parsed_date >= start_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "FMI"})
    except Exception as e:
        pass
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_country_reports_elibrary(start_date_str, end_date_str):
    """Extractor FMI - Country Reports (Bypass de Tapestry 5 AJAX Lazy-Loading)"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
    }
    
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    base_domain = "https://www.elibrary.imf.org"
    url_overview = f"{base_domain}/view/journals/002/002-overview.xml"
    
    try:
        # FASE 1: Extraer los tokens dinámicos de AJAX para los años recientes
        res = requests.get(url_overview, headers=headers, timeout=15)
        if res.status_code != 200: return pd.DataFrame()
        
        soup = BeautifulSoup(res.text, 'html.parser')
        
        ajax_links = []
        current_year = datetime.datetime.now().year
        # Buscamos los enlaces de expansión para el año actual y el anterior
        target_years = [str(current_year), str(current_year - 1)] 
        
        for li in soup.find_all('div', attrs={'data-toc-role': 'li'}):
            label_div = li.find('div', class_='label')
            if not label_div: continue
            
            texto_label = label_div.get_text()
            if any(year in texto_label for year in target_years):
                a_tag = li.find('a', class_='ajax-control')
                if a_tag and a_tag.has_attr('href'):
                    ajax_links.append(base_domain + a_tag['href'])
        
        # FASE 2: Interceptar y "deshidratar" las respuestas AJAX de Tapestry
        headers_ajax = headers.copy()
        headers_ajax['X-Requested-With'] = 'XMLHttpRequest' # Engañamos al framework
        headers_ajax['Accept'] = 'application/json, text/javascript, */*; q=0.01'
        
        for ajax_url in ajax_links:
            try:
                res_ajax = requests.get(ajax_url, headers=headers_ajax, timeout=15)
                if res_ajax.status_code != 200: continue
                
                data = res_ajax.json()
                
                # Extraemos el HTML inyectado dentro del nodo "zones"
                html_fragment = ""
                if "zones" in data:
                    for zone_id, html_content in data["zones"].items():
                        html_fragment += html_content
                        
                if not html_fragment: continue
                
                # FASE 3: Parsear el HTML revelado
                soup_fragment = BeautifulSoup(html_fragment, 'html.parser')
                
                for a_tag in soup_fragment.find_all('a', href=True):
                    href = a_tag['href']
                    titulo = a_tag.get_text(strip=True)
                    
                    # Filtro de sanidad: debe ser un artículo real
                    if '/view/journals/002/' in href and len(titulo) > 15:
                        link_real = base_domain + href if href.startswith('/') else href
                        
                        # Buscamos la fecha subiendo hasta 3 niveles en el DOM
                        date_str = ""
                        for padre in a_tag.find_parents(['div', 'li'], limit=3):
                            texto_padre = padre.get_text(separator=" ", strip=True)
                            
                            # Caza fechas en formatos "Mar 05, 2026" o "05 March 2026"
                            match = re.search(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2}?,?\s*\d{4}', texto_padre)
                            if not match:
                                match = re.search(r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}', texto_padre)
                                
                            if match:
                                date_str = match.group(0)
                                break # Encontramos la fecha, salimos del bucle
                                
                        parsed_date = None
                        if date_str:
                            try:
                                parsed_date = parser.parse(date_str)
                                if parsed_date.tzinfo is not None: parsed_date = parsed_date.replace(tzinfo=None)
                            except: pass
                            
                        # Evaluación final
                        if parsed_date and parsed_date >= start_date:
                            if not any(r['Link'] == link_real for r in rows):
                                rows.append({"Date": parsed_date, "Title": titulo, "Link": link_real, "Organismo": "FMI"})
            except:
                continue # Aislamiento de fallos
                
    except Exception as e:
        pass
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_pub_inst_fmi(start_date_str, end_date_str):
    """Extractor FMI - Vía directa por API Next.js (El Regalo)"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'application/json, text/plain, */*'
    }
    
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    # 1. CAZADOR DE BUILD ID (Para que tu código no caduque nunca)
    build_id = "OPXKbpp2La91iW-gTVkBX" # Tu regalo como plan de respaldo
    try:
        res_html = requests.get("https://www.imf.org/en/publications", headers=headers, timeout=15)
        # Buscamos el código dinámico oculto en la página principal
        match = re.search(r'"buildId":"([^"]+)"', res_html.text)
        if match:
            build_id = match.group(1)
    except:
        pass

    # 2. CONSTRUCCIÓN DE LOS ENLACES JSON DIRECTOS
    endpoints_json = [
        f"https://www.imf.org/_next/data/{build_id}/en/publications/fm.json",
        f"https://www.imf.org/_next/data/{build_id}/en/publications/weo.json",
        f"https://www.imf.org/_next/data/{build_id}/en/publications/gfsr.json"
    ]
    
    for url in endpoints_json:
        try:
            # Ahora pedimos el JSON limpio, evadiendo el HTML
            res = requests.get(url, headers=headers, timeout=15)
            if res.status_code != 200: continue
            data = res.json()
            
            # Buscador recursivo dentro del JSON
            def extraer_issues(obj):
                if isinstance(obj, dict):
                    if "issuePage" in obj and isinstance(obj["issuePage"], dict) and "results" in obj["issuePage"]:
                        for r in obj["issuePage"]["results"]: yield r
                    for k, v in obj.items(): yield from extraer_issues(v)
                elif isinstance(obj, list):
                    for item in obj: yield from extraer_issues(item)

            for issue in extraer_issues(data):
                titulo = issue.get("title", {}).get("jsonValue", {}).get("value", "")
                link_raw = issue.get("url", {}).get("url", "") or issue.get("url", {}).get("path", "")
                if not titulo or not link_raw: continue
                
                link_real = link_raw if link_raw.startswith("http") else "https://www.imf.org" + link_raw
                
                d_str = issue.get("publicationDate", {}).get("jsonValue", {}).get("value", "")
                if d_str:
                    try:
                        parsed_date = parser.parse(d_str)
                        if parsed_date.tzinfo is not None: parsed_date = parsed_date.replace(tzinfo=None)
                        if parsed_date >= start_date and not any(r['Link'] == link_real for r in rows):
                            rows.append({"Date": parsed_date, "Title": titulo, "Link": link_real, "Organismo": "FMI"})
                    except: pass
        except:
            continue
            
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

## FMI - Publicaciones Institucionales - F&D Magazine
@st.cache_data(show_spinner=False)
def load_pub_inst_fandd(start_date_str, end_date_str):
    """
    Extrae ediciones completas de la revista F&D Magazine del FMI
    """
    import requests
    import json
    import re
    import datetime
    import pandas as pd
    from dateutil import parser

    print("="*50)
    print("📘 CARGANDO F&D MAGAZINE")
    print(f"   Fechas: {start_date_str} a {end_date_str}")
    print("="*50)

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"   Rango parseado: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"   ⚠️ Error en fechas, usando rango por defecto")

    url = "https://www.imf.org/en/publications/fandd/issues"
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
    }

    rows = []

    try:
        print(f"   📡 Solicitando: {url}")
        res = requests.get(url, headers=headers, timeout=15)
        print(f"   Status Code: {res.status_code}")
        
        if res.status_code != 200:
            print(f"   ❌ Error al acceder: {res.status_code}")
            return pd.DataFrame()

        # Buscar el JSON de Next.js
        match = re.search(r'<script id="__NEXT_DATA__" type="application/json">(.*?)</script>', res.text, re.DOTALL)
        
        if not match:
            print("   ❌ No se encontró __NEXT_DATA__")
            return pd.DataFrame()

        data = json.loads(match.group(1))
        print("   ✅ JSON encontrado")
        
        # Buscar los issues
        results = []
        
        # Ruta según el HTML que proporcionaste
        try:
            page_props = data.get('props', {}).get('pageProps', {})
            component_props = page_props.get('componentProps', {})
            
            for comp_id, comp_data in component_props.items():
                if isinstance(comp_data, dict) and 'issueList' in comp_data:
                    issue_list = comp_data['issueList']
                    if isinstance(issue_list, dict) and 'results' in issue_list:
                        results = issue_list['results']
                        print(f"   ✅ Encontrados {len(results)} issues en componente: {comp_id}")
                        break
        except Exception as e:
            print(f"   ⚠️ Error navegando: {e}")
        
        if not results:
            print("   ❌ No se encontraron issues")
            return pd.DataFrame()

        meses_map = {
            'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
            'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12,
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
        }

        for issue in results:
            issue_title = issue.get('issueTitle', {}).get('jsonValue', {}).get('value', '')
            issue_label = issue.get('issueLabel', {}).get('jsonValue', {}).get('value', '')
            issue_url = issue.get('url', {}).get('url', '')
            
            if not issue_url and issue.get('url', {}).get('path'):
                issue_url = "https://www.imf.org" + issue.get('url', {}).get('path', '')
            
            fecha_texto = issue_label if issue_label else issue_title
            
            match_date = re.search(r'([A-Za-z]+)\s+(\d{4})', fecha_texto, re.IGNORECASE)
            if not match_date:
                print(f"   ⚠️ No se pudo parsear fecha: '{fecha_texto}'")
                continue
            
            mes_str = match_date.group(1).lower()
            año = int(match_date.group(2))
            mes_num = meses_map.get(mes_str, 1)
            
            issue_date = datetime.datetime(año, mes_num, 15)
            
            if issue_date < start_date or issue_date > end_date:
                print(f"   ⏭️ Fuera de rango: {issue_date.strftime('%Y-%m')} - {issue_title}")
                continue
            
            title_clean = re.sub(r'\s+', ' ', issue_title).strip()
            if not title_clean:
                title_clean = fecha_texto
            
            titulo_final = f"F&D: {issue_label} - {title_clean}" if issue_label else f"F&D: {title_clean}"
            
            rows.append({
                "Date": issue_date,
                "Title": titulo_final,
                "Link": issue_url if issue_url else f"https://www.imf.org/en/publications/fandd/issues/{año}/{mes_num:02d}",
                "Organismo": "FMI"
            })
            print(f"   ✅ AGREGADO: {issue_date.strftime('%Y-%m-%d')} - {titulo_final[:60]}...")
        
    except Exception as e:
        print(f"   ❌ Error general: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
    
    print(f"   📊 TOTAL F&D: {len(df)} ediciones")
    print("="*50)
    return df

@st.cache_data(show_spinner=False)
def load_pub_inst_bm(start_date_str, end_date_str):
    """Extractor para Publicaciones Institucionales (Colecciones Específicas) del BM"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}
    
    # IDs exactos de las 3 colecciones
    scopes = [
        '4c48a649-7773-4d0f-b441-f5fc7e8d67f8', # Business Ready
        '09c5e8fc-187f-5c2f-a077-3e03044c7b62', # Perspectivas económicas mundiales
        '3d9bbbf6-c007-5043-b655-04d8a1cfbfb2'  # Tercera colección
    ]
    
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    # Iteramos sobre cada una de las colecciones
    for scope in scopes:
        page = 0
        while True:
            try:
                # Al pasarle el 'scope', la API restringe la búsqueda SOLO a esa colección
                params = {
                    'scope': scope,
                    'sort': 'dc.date.issued,DESC', 
                    'page': page, 
                    'size': 20
                }
                res = requests.get(base_url, headers=headers, params=params, timeout=15)
                data = res.json()
                
                objects = data.get('_embedded', {}).get('searchResult', {}).get('_embedded', {}).get('objects', [])
                if not objects: break
                
                items_found = 0
                for obj in objects:
                    item = obj.get('_embedded', {}).get('indexableObject', {})
                    meta = item.get('metadata', {})
                    
                    title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                    date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                    
                    parsed_date = None
                    if date_s:
                        try: parsed_date = parser.parse(date_s)
                        except: pass
                    
                    if not parsed_date or parsed_date < start_date: continue
                    
                    link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '')
                    if not link: link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": title, "Link": link, "Organismo": "BM"})
                        items_found += 1
                
                if items_found == 0: break
                page += 1
                if page > 3: break # Límite de seguridad
                time.sleep(0.2)
            except:
                break
                
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df



    # --- SECCIÓN: INVESTIGACIÓN ---
## - Working Papers - FMI
@st.cache_data(show_spinner=False)
def load_working_papers_fmi(start_date_str, end_date_str):
    """
    Extractor FMI - Working Papers usando Crossref API
    Los Working Papers del FMI tienen DOIs con prefix 10.5089
    """
    import requests
    import datetime
    import re
    from dateutil import parser
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 FMI Working Papers (Crossref API): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    
    # API de Crossref
    url = "https://api.crossref.org/works"
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json"
    }
    
    # Parámetros de búsqueda
    params = {
        "filter": f"from-pub-date:{start_date.strftime('%Y-%m-%d')},until-pub-date:{end_date.strftime('%Y-%m-%d')},prefix:10.5089",
        "rows": 100,
        "sort": "published-online",
        "order": "desc"
    }
    
    try:
        print(f"📡 Solicitando DOIs del FMI a Crossref API...")
        
        response = requests.get(url, headers=headers, params=params, timeout=30, verify=False)
        
        print(f"   Status Code: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            items = data.get('message', {}).get('items', [])
            
            print(f"📚 Documentos encontrados en Crossref: {len(items)}")
            
            for item in items:
                # Extraer título
                titulo = item.get('title', [''])[0] if item.get('title') else ''
                if not titulo:
                    continue
                
                # Extraer DOI
                doi = item.get('DOI', '')
                if not doi:
                    continue
                
                # Construir URL del DOI
                link = f"https://doi.org/{doi}"
                
                # Extraer fecha
                pub_date = item.get('published-print', {}) or item.get('published-online', {})
                date_parts = pub_date.get('date-parts', [[]])[0]
                
                parsed_date = None
                if len(date_parts) >= 3:
                    try:
                        parsed_date = datetime.datetime(date_parts[0], date_parts[1], date_parts[2])
                    except:
                        pass
                elif len(date_parts) >= 2:
                    try:
                        parsed_date = datetime.datetime(date_parts[0], date_parts[1], 1)
                    except:
                        pass
                
                if not parsed_date:
                    continue
                
                # Verificar que sea Working Paper
                container = item.get('container-title', [''])[0] if item.get('container-title') else ''
                is_working_paper = 'working paper' in container.lower() or 'imf working' in container.lower()
                
                if not is_working_paper:
                    is_working_paper = 'working paper' in titulo.lower()
                
                if not is_working_paper:
                    continue
                
                if "coming soon" in titulo.lower():
                    continue
                
                if start_date <= parsed_date <= end_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "FMI"
                        })
                        print(f"   ✅ {parsed_date.strftime('%Y-%m-%d')}: {titulo[:60]}...")
                        
        else:
            print(f"❌ Error en Crossref API: {response.status_code}")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 FMI Working Papers - Total final: {len(df)}")
    return df

# ========== INVESTIGACIÓN CEMLA (Latin American Journal of Central Banking) ==========
@st.cache_data(show_spinner=False)
def load_investigacion_cemla(start_date_str, end_date_str):
    """
    Extractor CEMLA - Latin American Journal of Central Banking
    Extrae fecha COMPLETA (con día) si está disponible en Crossref
    """
    import requests
    import datetime
    from dateutil import parser
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    print("="*60)
    print("🔍 CEMLA INVESTIGACIÓN - Buscando fechas completas (con día)")
    print("="*60)
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 Rango: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    issn = "2666-1438"
    base_url = "https://api.crossref.org/works"
    
    # Buscar mes por mes para tener mejor control
    current = start_date.replace(day=1)
    
    while current <= end_date:
        year = current.year
        month = current.month
        
        # Último día del mes
        if month == 12:
            last_day = 31
        elif month in [4, 6, 9, 11]:
            last_day = 30
        else:
            last_day = 28 if year % 4 != 0 else 29
        
        fecha_inicio = f"{year}-{month:02d}-01"
        fecha_fin = f"{year}-{month:02d}-{last_day}"
        
        print(f"\n📆 Buscando {year}-{month:02d}...")
        
        params = {
            "filter": f"from-pub-date:{fecha_inicio},until-pub-date:{fecha_fin},issn:{issn}",
            "rows": 50,
            "sort": "published-online",
            "order": "desc"
        }
        
        try:
            response = requests.get(base_url, params=params, timeout=30, verify=False)
            
            if response.status_code == 200:
                data = response.json()
                items = data.get('message', {}).get('items', [])
                
                if items:
                    print(f"   📚 Artículos: {len(items)}")
                    
                    for item in items:
                        titulo = item.get('title', [''])[0] if item.get('title') else ''
                        doi = item.get('DOI', '')
                        link = f"https://doi.org/{doi}" if doi else ''
                        
                        if not titulo or not link:
                            continue
                        
                        # ========== INTENTAR OBTENER FECHA COMPLETA ==========
                        fecha_completa = None
                        
                        # 1. Probar con 'published-online' (puede tener día)
                        pub_online = item.get('published-online', {})
                        if pub_online:
                            date_parts = pub_online.get('date-parts', [[]])[0]
                            if len(date_parts) >= 3:
                                try:
                                    fecha_completa = datetime.datetime(date_parts[0], date_parts[1], date_parts[2])
                                    print(f"      📅 Online: {fecha_completa.strftime('%Y-%m-%d')}")
                                except:
                                    pass
                        
                        # 2. Probar con 'issued' (fecha de publicación)
                        if not fecha_completa:
                            issued = item.get('issued', {})
                            if issued:
                                date_parts = issued.get('date-parts', [[]])[0]
                                if len(date_parts) >= 3:
                                    try:
                                        fecha_completa = datetime.datetime(date_parts[0], date_parts[1], date_parts[2])
                                        print(f"      📅 Issued: {fecha_completa.strftime('%Y-%m-%d')}")
                                    except:
                                        pass
                        
                        # 3. Probar con 'posted-online'
                        if not fecha_completa:
                            posted = item.get('posted-online', {})
                            if posted:
                                date_parts = posted.get('date-parts', [[]])[0]
                                if len(date_parts) >= 3:
                                    try:
                                        fecha_completa = datetime.datetime(date_parts[0], date_parts[1], date_parts[2])
                                        print(f"      📅 Posted: {fecha_completa.strftime('%Y-%m-%d')}")
                                    except:
                                        pass
                        
                        # 4. Fallback: usar el primer día del mes (si no hay día)
                        if not fecha_completa:
                            fecha_completa = datetime.datetime(year, month, 1)
                            print(f"      ⚠️ Fallback: {fecha_completa.strftime('%Y-%m-%d')} (sin día específico)")
                        
                        # Filtrar por rango
                        if start_date <= fecha_completa <= end_date:
                            rows.append({
                                "Date": fecha_completa,
                                "Title": titulo,
                                "Link": link,
                                "Organismo": "CEMLA"
                            })
                            print(f"      ✅ AGREGADO: {fecha_completa.strftime('%Y-%m-%d')}")
                        else:
                            print(f"      ⏭️ Fuera de rango: {fecha_completa.strftime('%Y-%m-%d')}")
                            
                else:
                    print(f"   📭 Sin artículos")
                    
            else:
                print(f"   ❌ Error: {response.status_code}")
                
        except Exception as e:
            print(f"   ❌ Error: {e}")
        
        # Siguiente mes
        if current.month == 12:
            current = current.replace(year=current.year + 1, month=1)
        else:
            current = current.replace(month=current.month + 1)
        
        time.sleep(0.5)
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"\n{'='*60}")
    print(f"📊 CEMLA Investigación - Total: {len(df)} documentos")
    if not df.empty:
        print("\n📅 Primeros 5 documentos con sus fechas:")
        for i, row in df.head(5).iterrows():
            print(f"   {row['Date'].strftime('%Y-%m-%d')}: {row['Title'][:60]}...")
    print(f"{'='*60}")
    
    return df

@st.cache_data(show_spinner=False)
def load_investigacion_fmi(start_date_str, end_date_str):
    """Extractor FMI - Blogs de Investigación (Vía Coveo API) - Versión mejorada"""
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 FMI Blogs - Rango solicitado: {start_date.date()} a {end_date.date()}")
    except Exception as e:
        print(f"⚠️ Error en fechas: {e}")
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"📅 Usando rango por defecto: {start_date.date()} a {end_date.date()}")

    rows = []
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",
        "Referer": "https://www.imf.org/",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    payload = {
        "aq": "@imftype==\"IMF Blog Page\" AND @syslanguage==\"English\"",
        "numberOfResults": 250,  # Aumentado para capturar más
        "sortCriteria": "@imfdate descending"
    }

    try:
        print("📡 Solicitando blogs a la API de Coveo...")
        res = requests.post(url, headers=headers, json=payload, timeout=15, verify=False)
        
        if res.status_code == 200:
            data = res.json()
            total_api = data.get('totalCount', 0)
            print(f"✅ Total de blogs en la API: {total_api}")
            
            documentos_filtrados = 0
            for item in data.get("results", []):
                titulo = item.get("title", "").strip()
                link = item.get("clickUri", "")
                
                # === MEJORA: Extraer fecha de múltiples formatos ===
                parsed_date = None
                raw_data = item.get("raw", {})
                
                # Formato 1: timestamp en milisegundos (el más común)
                raw_date = raw_data.get("date")
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except:
                        pass
                
                # Formato 2: fecha como string ISO
                if not parsed_date:
                    date_str = raw_data.get("date") or raw_data.get("publisheddate") or raw_data.get("publicationdate")
                    if date_str and isinstance(date_str, str):
                        try:
                            parsed_date = parser.parse(date_str)
                        except:
                            pass
                
                # Formato 3: intentar con cualquier campo que parezca fecha
                if not parsed_date:
                    for key in ['date', 'publisheddate', 'publicationdate', 'createddate', 'lastmodified']:
                        val = raw_data.get(key)
                        if val:
                            try:
                                if isinstance(val, (int, float)):
                                    parsed_date = datetime.datetime.fromtimestamp(val / 1000.0)
                                elif isinstance(val, str):
                                    parsed_date = parser.parse(val)
                                if parsed_date:
                                    break
                            except:
                                continue
                
                if not titulo or not link or not parsed_date:
                    continue
                
                # Depuración: mostrar las fechas que se están procesando
                print(f"   📅 Procesando: {parsed_date.strftime('%Y-%m-%d')} - {titulo[:50]}...")
                
                # Filtrar por el rango de fechas
                if start_date <= parsed_date <= end_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date, 
                            "Title": titulo, 
                            "Link": link, 
                            "Organismo": "FMI"
                        })
                        documentos_filtrados += 1
                        print(f"      ✅ AGREGADO: {parsed_date.strftime('%Y-%m-%d')}")
            
            print(f"\n📊 Total de blogs en el rango {start_date.date()} a {end_date.date()}: {documentos_filtrados}")
            
        else:
            print(f"❌ Error en la API: {res.status_code}")
            
    except Exception as e:
        print(f"❌ Error en load_investigacion_fmi: {e}")
        import traceback
        traceback.print_exc()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    return df

@st.cache_data(show_spinner=False)
def load_investigacion_bm(start_date_str, end_date_str):
    """Extractor para Investigación del BM (Filtra y excluye los que son 'Reports')"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}
    
    # ID exacto de la comunidad de Investigación
    scope_id = '06251f8a-62c2-59fb-add5-ec0993fc20d9'
    
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows, page = [], 0
    while True:
        try:
            params = {
                'scope': scope_id, 
                'sort': 'dc.date.issued,DESC', 
                'page': page, 
                'size': 20
            }
            res = requests.get(base_url, headers=headers, params=params, timeout=15)
            data = res.json()
            
            objects = data.get('_embedded', {}).get('searchResult', {}).get('_embedded', {}).get('objects', [])
            if not objects: break
            
            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})
                
                # Extraer Título y Fecha
                title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                
                parsed_date = None
                if date_s:
                    try: parsed_date = parser.parse(date_s)
                    except: pass
                
                if not parsed_date or parsed_date < start_date: continue
                
                # --- NUEVO FILTRO ANTI-REPORTES ---
                # Buscamos en el abstract o en la descripción general
                abstract_list = meta.get('dc.description.abstract', [])
                desc_list = meta.get('dc.description', [])
                
                description = ""
                if abstract_list: description = abstract_list[0].get('value', '').lower()
                elif desc_list: description = desc_list[0].get('value', '').lower()
                
                # Si la palabra exacta "report" está en la descripción, lo saltamos
                # Usamos \b para que sea la palabra exacta y no algo como "reporting"
                if re.search(r'\breport\b', description):
                    continue
                # ----------------------------------
                
                # Link permanente
                link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '')
                if not link: link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"
                
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": title, "Link": link, "Organismo": "BM"})
                    items_found += 1
            
            if items_found == 0: break
            page += 1
            if page > 3: break # Límite para evitar búsquedas infinitas
            time.sleep(0.2)
        except:
            break
            
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

## OCDE - INVESTIGACION

@st.cache_data(show_spinner=False)
def load_investigacion_ocde(start_date_str, end_date_str):
    """Extractor OCDE - Working Papers (API oficial con paginación)"""
    import requests
    import datetime
    import re
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 OCDE Investigación: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    # API base de la OCDE
    base_url = "https://api.oecd.org/webcms/search/faceted-search"
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json"
    }
    
    page = 0
    page_size = 50  # Número de resultados por página
    max_pages = 10  # Límite de seguridad (500 documentos máximo)
    documentos_procesados = 0
    
    print("📡 Solicitando Working Papers a la API de la OCDE (con paginación)...")
    
    try:
        while page < max_pages:
            params = {
                "siteName": "oecd",
                "interfaceLanguage": "en",
                "orderBy": "mostRecent",
                "pageSize": page_size,
                "page": page,
                "facets": "oecd-languages:en",
                "hiddenFacets": "oecd-content-types:publications/working-papers"
            }
            
            print(f"   📄 Procesando página {page + 1}...")
            response = requests.get(base_url, params=params, headers=headers, timeout=15)
            
            if response.status_code != 200:
                print(f"   ❌ Error en página {page + 1}: {response.status_code}")
                break
            
            data = response.json()
            
            # Buscar los resultados
            results = []
            if "results" in data:
                results = data["results"]
            else:
                print(f"   ⚠️ Estructura inesperada en página {page + 1}")
                break
            
            if not results:
                print(f"   📭 No hay más resultados en página {page + 1}")
                break
            
            # Contar cuántos documentos del mes encontramos en esta página
            documentos_en_pagina = 0
            fecha_mas_reciente = None
            fecha_mas_antigua = None
            
            for item in results:
                titulo = item.get("title", "") or item.get("name", "")
                link = item.get("url", "") or item.get("link", "")
                
                if not titulo or not link:
                    continue
                
                # Extraer fecha del campo publicationDateTime
                fecha_texto = item.get("publicationDateTime", "")
                
                parsed_date = None
                if fecha_texto:
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        continue
                
                if not parsed_date:
                    continue
                
                # Actualizar fechas extremas
                if fecha_mas_reciente is None or parsed_date > fecha_mas_reciente:
                    fecha_mas_reciente = parsed_date
                if fecha_mas_antigua is None or parsed_date < fecha_mas_antigua:
                    fecha_mas_antigua = parsed_date
                
                # Si el documento es más antiguo que start_date, podemos parar porque
                # los resultados están ordenados por fecha descendente
                if parsed_date < start_date:
                    # Ya no hay más documentos del mes en esta página ni en las siguientes
                    print(f"   ⏹️ Documento más antiguo que {start_date.strftime('%Y-%m')}, deteniendo paginación")
                    # Salimos del while principal
                    page = max_pages
                    break
                
                # Filtrar por rango de fechas
                if parsed_date >= start_date and parsed_date <= end_date:
                    # Limpiar título
                    titulo = re.sub(r'\s+', ' ', titulo).strip()
                    
                    # Asegurar URL absoluta
                    if link.startswith('/'):
                        link = f"https://www.oecd.org{link}"
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "OCDE"
                    })
                    documentos_en_pagina += 1
                    documentos_procesados += 1
            
            print(f"   📊 Página {page + 1}: {documentos_en_pagina} documentos en el rango")
            
            # Si no encontramos documentos en esta página y ya pasamos la fecha límite, paramos
            if documentos_en_pagina == 0 and fecha_mas_antigua and fecha_mas_antigua < start_date:
                print(f"   ⏹️ Fin de resultados para el mes solicitado")
                break
            
            # Si encontramos menos de page_size documentos, probablemente es la última página
            if len(results) < page_size:
                print(f"   📭 Última página alcanzada")
                break
            
            page += 1
            # Pequeña pausa para no sobrecargar la API
            time.sleep(0.3)
        
        print(f"\n📊 Total documentos OCDE encontrados: {documentos_procesados}")
        
    except Exception as e:
        print(f"❌ Error en load_investigacion_ocde: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 OCDE Investigación - Total final: {len(df)}")
    return df

# --- SECCIÓN: DISCURSOS ---
## -- Banco de Inglaterra -- Bank of England (BoE)
@st.cache_data(show_spinner=False)
def load_discursos_boe(start_date_str, end_date_str):
    """Extractor Automático BoE - Vía RSS con formato consistente 'Autor: Título'"""
    import requests
    from bs4 import BeautifulSoup
    import pandas as pd
    import datetime
    import re
    from dateutil import parser

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    url = "https://www.bankofengland.co.uk/rss/speeches"
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []

    def extract_author_from_title(title):
        """Extrae el nombre del autor del título en varios formatos"""
        autor = ""
        titulo_limpio = title
        
        # Patrón 1: "Título − speech by Autor" (con guión largo o corto)
        match = re.search(r'(?i)\s*[\-–—]\s*speech\s+by\s+(.+?)$', title)
        if match:
            autor = clean_author_name(match.group(1).strip())
            # Eliminar TODO desde el guión hasta el final
            titulo_limpio = re.sub(r'(?i)\s*[\-–—]\s*speech\s+by\s+.*$', '', title).strip()
            return autor, titulo_limpio
        
        # Patrón 2: "Speech by Autor: Título" o "Speech by Autor - Título"
        match = re.search(r'(?i)^speech\s+by\s+([^:—-]+)[:—-]\s*(.+)$', title)
        if match:
            autor = clean_author_name(match.group(1).strip())
            titulo_limpio = match.group(2).strip()
            return autor, titulo_limpio
        
        # Patrón 3: "Autor: Título" (ya está bien formateado)
        match = re.search(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*:\s*(.+)$', title)
        if match:
            autor = clean_author_name(match.group(1))
            titulo_limpio = match.group(2)
            return autor, titulo_limpio
        
        # Patrón 4: "Título by Autor" (sin "speech")
        match = re.search(r'(?i)\s+by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)$', title)
        if match:
            autor = clean_author_name(match.group(1))
            titulo_limpio = re.sub(r'(?i)\s+by\s+.*$', '', title).strip()
            return autor, titulo_limpio
        
        return None, title

    try:
        res = requests.get(url, headers=headers, timeout=10)
        if res.status_code == 200:
            soup = BeautifulSoup(res.content, "xml")
            items = soup.find_all("item")

            for item in items:
                titulo_raw = item.find("title").text if item.find("title") else ""
                link = item.find("link").text if item.find("link") else ""
                fecha_raw = item.find("pubDate").text if item.find("pubDate") else ""

                if not titulo_raw or not link or not fecha_raw:
                    continue

                try:
                    parsed_date = parser.parse(fecha_raw)
                    if parsed_date.tzinfo is not None:
                        parsed_date = parsed_date.replace(tzinfo=None)
                except:
                    continue

                if start_date <= parsed_date <= end_date:
                    # Extraer autor y título limpio
                    autor, titulo_limpio = extract_author_from_title(titulo_raw)
                    
                    # LIMPIEZA DIRECTA: eliminar específicamente " − speech" o " −" al final
                    # Primero, eliminar " − speech" (con el guión especial)
                    titulo_limpio = titulo_limpio.replace(' − speech', '').replace(' - speech', '').replace('— speech', '')
                    # Luego, eliminar " −" solitario al final
                    titulo_limpio = titulo_limpio.replace(' −', '').replace(' -', '').replace('—', '')
                    # Eliminar espacios sobrantes al final
                    titulo_limpio = titulo_limpio.rstrip()
                    
                    # Construir título final en formato "Autor: Título"
                    if autor:
                        titulo_final = f"{autor}: {titulo_limpio}"
                    else:
                        titulo_final = titulo_limpio
                    
                    # Limpieza final de espacios múltiples
                    titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                    titulo_final = titulo_final.strip('"').strip("'").strip()
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo_final,
                            "Link": link,
                            "Organismo": "BoE (Inglaterra)"
                        })
    except Exception as e:
        print(f"Error en load_discursos_boe: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values(by="Date", ascending=False)
    return df


## -- FMI - Discursos 
@st.cache_data(show_spinner=False)
def load_discursos_fmi(start_date_str, end_date_str):
    """Extractor FMI - Discursos y Transcripts (Coveo API + Scraping Blindado)"""
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",
        "Referer": "https://www.imf.org/",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    payload = {
        "aq": "@imftype==(\"Speech\",\"Transcript\") AND @syslanguage==\"English\"",
        "numberOfResults": 150,
        "sortCriteria": "@imfdate descending"
    }

    try:
        res = requests.post(url, headers=headers, json=payload, timeout=15)
        if res.status_code == 200:
            data = res.json()
            for item in data.get("results", []):
                titulo_raw = item.get("title", "").strip()
                link = item.get("clickUri", "")
                raw_data = item.get("raw", {})
                raw_date = raw_data.get("date")

                parsed_date = None
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(
                            raw_date / 1000.0)
                    except:
                        pass
                if not titulo_raw or not link or not parsed_date:
                    continue

                # SOLO procesamos el documento si está en el rango de fechas para ahorrar tiempo
                if start_date <= parsed_date <= end_date:
                    if not any(r['Link'] == link for r in rows):

                        # 1. Autor oficial de la etiqueta API
                        autor = raw_data.get("imfspeaker", "")
                        if isinstance(autor, list) and len(autor) > 0:
                            autor = autor[0]

                        # 2. CAZADOR BLINDADO (Si la API viene vacía, visitamos el link)
                        if not autor:
                            try:
                                link_req = link if link.startswith(
                                    'http') else "https://www.imf.org" + link
                                art_res = requests.get(
                                    link_req, headers=headers, timeout=10)

                                # TRITURADORA: Reemplazamos cualquier etiqueta HTML (<p>, <em>, <strong>) por un espacio
                                art_text = re.sub(
                                    r'<[^>]+>', ' ', art_res.text)
                                # Colapsamos múltiples espacios y saltos de línea en uno solo
                                art_text = re.sub(r'\s+', ' ', art_text)

                                # Buscamos "Speakers: Nombre" deteniéndonos en la coma o en su cargo
                                match_speaker = re.search(
                                    r'(?:Speakers?|Participants?)\s*:\s*([A-ZÀ-ÿ][A-Za-zÀ-ÿ\s\.\'-]{3,40}?)\s*(?:,|-|–|—|Director|Managing|Deputy|Senior|Head|Minister|Secretary)', art_text, re.IGNORECASE)
                                if match_speaker:
                                    autor = match_speaker.group(1).strip()
                            except Exception as e:
                                pass

                        # Limpiamos el nombre encontrado y cortamos si hay dos personas (ej: "Kristalina and Nigel")
                        if autor:
                            autor = re.split(
                                r'\s+and\s+|\s*&\s*', autor, flags=re.IGNORECASE)[0].strip()
                            autor = clean_author_name(autor)

                        # =========================================================
                        # LIMPIEZA ESTÉTICA
                        # =========================================================
                        titulo_limpio = titulo_raw

                        # Quitar sufijos (Ej: - Keynote Speech)
                        patron_sufijo = re.compile(
                            r"(?i)\s*[\-–—]\s*.*?\b(speech|remarks|statement|address|transcript|keynote)\b\s+by\s+.*$")
                        titulo_limpio = patron_sufijo.sub(
                            "", titulo_limpio).strip().strip('"').strip("'").strip()

                        # Quitar prefijos de Transcripts (Ej: Press Briefing Transcript:)
                        patron_prefijo = re.compile(
                            r"(?i)^(Press Briefing Transcript|Transcript of Press Briefing|Transcript)\s*[:\-]\s*")
                        titulo_limpio = patron_prefijo.sub(
                            "", titulo_limpio).strip()

                        # Transformar "Remarks by Autor, ..."
                        patron_remarks = re.compile(
                            r"(?i)^(remarks|speech|statement|address)\s+by\s+([^,:]+)[,\-]?\s*")
                        match_remarks = patron_remarks.match(titulo_limpio)
                        if match_remarks:
                            autor_detectado = clean_author_name(
                                match_remarks.group(2))
                            if not autor:
                                autor = autor_detectado
                            titulo_limpio = patron_remarks.sub(
                                f"{autor}: ", titulo_limpio)

                        # Inyectar Autor Final
                        if autor:
                            if titulo_limpio.lower().startswith(f"{autor.lower()},"):
                                titulo_final = re.sub(
                                    rf"(?i)^{re.escape(autor)},", f"{autor}:", titulo_limpio)
                            elif titulo_limpio.lower().startswith(f"{autor.lower()}:"):
                                titulo_final = titulo_limpio
                            elif titulo_limpio.lower().startswith(autor.lower()):
                                titulo_final = re.sub(
                                    rf"(?i)^{re.escape(autor)}\s*", f"{autor}: ", titulo_limpio)
                            elif autor.lower() not in titulo_limpio.lower():
                                titulo_final = f"{autor}: {titulo_limpio}"
                            else:
                                titulo_final = titulo_limpio
                        else:
                            titulo_final = titulo_limpio
                        # =========================================================

                        rows.append(
                            {"Date": parsed_date, "Title": titulo_final, "Link": link, "Organismo": "FMI"})
    except:
        pass

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


# --- SECCIÓN: DISCURSOS ---
@st.cache_data(show_spinner=False)
def load_data_ecb(start_date_str, end_date_str):
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        anios_num = list(range(start_date.year, end_date.year + 1))
    except: anios_num = [2026, 2025, 2024]
    for year in anios_num:
        url = f"https://www.ecb.europa.eu/press/key/date/{year}/html/index.en.html"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            for a in soup.find_all('a', href=True):
                href = a['href']
                if f'/press/key/date/{year}/html/' in href and href.endswith('.html') and 'index' not in href:
                    link = "https://www.ecb.europa.eu" + href if href.startswith('/') else href
                    titulo_raw = a.get_text(strip=True)
                    if len(titulo_raw) < 5: continue
                    parent = a.find_parent(['dd', 'div', 'li'])
                    if not parent: continue
                    dt = parent.find_previous_sibling('dt')
                    fecha_str = dt.get_text(strip=True) if dt else ""
                    try: parsed_date = parser.parse(fecha_str)
                    except: continue
                    autor = ""
                    sub = parent.find('div', class_='subtitle')
                    if sub:
                        match = re.search(r'\b(?:by|with)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', sub.get_text(separator=' ', strip=True))
                        if match: autor = clean_author_name(match.group(1))
                    final_t = f"{autor}: {titulo_raw}" if autor and autor not in titulo_raw else titulo_raw
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": final_t, "Link": link, "Organismo": "ECB (Europa)"})
        except: pass
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_bis():
    urls = ["https://www.bis.org/api/document_lists/cbspeeches.json", "https://www.bis.org/api/document_lists/bcbs_speeches.json", "https://www.bis.org/api/document_lists/mgmtspeeches.json"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    for url in urls:
        try:
            response = requests.get(url, headers=headers, timeout=10)
            data = response.json()
            for path, speech in data.get("list", {}).items():
                title = html.unescape(speech.get("short_title", ""))
                date_str = speech.get("publication_start_date", "")
                link = "https://www.bis.org" + path + (".htm" if not path.endswith(".htm") else "")
                rows.append({"Date": date_str, "Title": title, "Link": link, "Organismo": "BPI"})
        except: continue
    df = pd.DataFrame(rows).drop_duplicates(subset=['Link']) if rows else pd.DataFrame()
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_bbk(start_date_str, end_date_str):
    base_url = "https://www.bundesbank.de/action/en/730564/bbksearch"
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows, page = [], 0
    while True:
        params = {'sort': 'bbksortdate desc', 'dateFrom': start_date_str, 'dateTo': end_date_str, 'pageNumString': str(page)}
        try: response = requests.get(base_url, headers=headers, params=params, timeout=10)
        except: break 
        soup = BeautifulSoup(response.text, 'html.parser')
        items = soup.find_all('li', class_='resultlist__item')
        if not items: break 
        for item in items:
            fecha_tag = item.find('span', class_='metadata__date')
            fecha_str = fecha_tag.text.strip() if fecha_tag else ""
            author_tag = item.find('span', class_='metadata__authors')
            author_str = clean_author_name(author_tag.text) if author_tag else ""
            data_div = item.find('div', class_='teasable__data')
            link, titulo = "", ""
            if data_div and data_div.find('a'):
                a_tag = data_div.find('a')
                link = "https://www.bundesbank.de" + a_tag.get('href', '') if a_tag.get('href', '').startswith('/') else a_tag.get('href', '')
                if a_tag.find('span', class_='link__label'): titulo = a_tag.find('span', class_='link__label').text.strip()
            if author_str and author_str not in titulo: titulo = f"{author_str}: {titulo}"
            if fecha_str and titulo: rows.append({"Date": fecha_str, "Title": titulo, "Link": link, "Organismo": "BBk (Alemania)"})
        if len(items) < 10: break
        page += 1
        time.sleep(0.3) 
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"], format='%d.%m.%Y', errors='coerce')
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_pboc(start_date_str, end_date_str):
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        url = "https://www.pbc.gov.cn/en/3688110/3688175/index.html" if page == 1 else f"https://www.pbc.gov.cn/en/3688110/3688175/0180081b-{page}.html"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            res.encoding = 'utf-8' 
            soup = BeautifulSoup(res.text, 'html.parser')
            items = soup.find_all('div', class_='ListR')
            if not items: break
            items_found = 0
            for item in items:
                date_span = item.find('span', class_='prhhdata')
                a_tag = item.find('a')
                if not date_span or not a_tag: continue
                try: parsed_date = parser.parse(date_span.get_text(strip=True))
                except: continue
                titulo_raw = html.unescape(a_tag.get('title', a_tag.get_text(strip=True)))
                link = "https://www.pbc.gov.cn" + a_tag.get('href', '') if a_tag.get('href', '').startswith('/') else a_tag.get('href', '')
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "PBoC (China)"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.5) 
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_fed(anios_num):
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    for year in anios_num:
        url = f"https://www.federalreserve.gov/newsevents/{year}-speeches.htm"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            if res.status_code == 404:
                url = "https://www.federalreserve.gov/newsevents/speeches.htm"
                res = requests.get(url, headers=headers, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            for a_tag in soup.find_all('a', href=True):
                if '/newsevents/speech/' in a_tag['href']:
                    link = "https://www.federalreserve.gov" + a_tag['href'] if a_tag['href'].startswith('/') else a_tag['href']
                    titulo = a_tag.get_text(strip=True)
                    parent = a_tag.find_parent('div', class_='row') or a_tag.parent
                    text = parent.get_text(separator=' | ', strip=True)
                    date_m = re.search(r'(\d{1,2}/\d{1,2}/\d{4}|\w+\s\d{1,2},\s\d{4})', text)
                    if date_m:
                        try:
                            parsed_date = parser.parse(date_m.group(1))
                            if parsed_date.year not in anios_num: continue
                            rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "Fed (Estados Unidos)"})
                        except: pass
        except: pass
    df = pd.DataFrame(rows).drop_duplicates(subset=['Link']) if rows else pd.DataFrame()
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_bdf(start_date_str, end_date_str):
    base_url = "https://www.banque-france.fr/en/governor-interventions"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 0
    while True:
        try:
            response = requests.get(base_url, headers=headers, params={'category[7052]': '7052', 'page': page}, timeout=12)
            soup = BeautifulSoup(response.text, 'html.parser')
            cards = soup.find_all('div', class_=lambda c: c and 'card' in c)
            if not cards: break
            items_found = 0
            for card in cards:
                a = card.find('a', href=True)
                if not a or not a.find('span', class_='title-truncation'): continue
                titulo_raw, link = a.find('span', class_='title-truncation').get_text(strip=True), "https://www.banque-france.fr" + a['href']
                date_s = card.find('small')
                if not date_s: continue
                fecha_clean = re.sub(r'(\d+)(st|nd|rd|th)\s+of\s+', r'\1 ', date_s.get_text(strip=True))
                try: parsed_date = parser.parse(fecha_clean)
                except: continue
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "BdF (Francia)"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_bm(start_date_str, end_date_str):
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 0
    while True:
        try:
            res = requests.get(base_url, headers=headers, params={'scope': 'b6a50016-276d-56d3-bbe5-891c8d18db24', 'sort': 'dc.date.issued,DESC', 'page': page, 'size': 20}, timeout=12)
            objects = res.json().get('_embedded', {}).get('searchResult', {}).get('_embedded', {}).get('objects', [])
            if not objects: break
            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})
                title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                try: parsed_date = parser.parse(date_s)
                except: continue
                link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '') or f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": title, "Link": link, "Organismo": "BM"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_boc(start_date_str, end_date_str):
    base_url = "https://www.bankofcanada.ca/press/speeches/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        try:
            res = requests.get(base_url, headers=headers, params={'mt_page': page}, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            articles = soup.find_all('div', class_=lambda c: c and ('mtt-result' in c or 'media' in c))
            if not articles: break
            items_found = 0
            for art in articles:
                h3 = art.find('h3', class_='media-heading')
                if not h3 or not h3.find('a'): continue
                titulo_raw, link = h3.find('a').text.strip(), h3.find('a')['href']
                date_s = art.find('span', class_='media-date')
                try: parsed_date = parser.parse(date_s.text.strip())
                except: continue
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "BoC (Canadá)"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_boj(start_date_str, end_date_str):
    base_url = "https://www.boj.or.jp/en/about/press/index.htm"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows = []
    try:
        response = requests.get(base_url, headers=headers, timeout=12)
        soup = BeautifulSoup(response.text, 'html.parser')
        table = soup.find('table', class_='js-tbl')
        if table:
            for tr in table.find('tbody').find_all('tr'):
                tds = tr.find_all('td')
                if len(tds) < 3: continue
                try: parsed_date = parser.parse(tds[0].get_text(strip=True).replace('\xa0', ' '))
                except: continue
                if parsed_date < start_date: continue
                a_tag = tds[2].find('a', href=True)
                if not a_tag: continue
                titulo_raw = a_tag.get_text(strip=True).strip('"')
                link = "https://www.boj.or.jp" + a_tag['href'] if a_tag['href'].startswith('/') else a_tag['href']
                rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "BoJ (Japón)"})
    except: pass
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_cef(start_date_str, end_date_str):
    base_url = "https://www.fsb.org/press/speeches-and-statements/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        url = f"{base_url}?dps_paged={page}"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            items = soup.find_all('div', class_='post-excerpt')
            if not items: break
            items_found = 0
            for item in items:
                title_tag = item.find('div', class_='post-title')
                if not title_tag or not title_tag.find('a'): continue
                a = title_tag.find('a')
                titulo_raw, link = a.get_text(strip=True), a['href']
                date_tag = item.find('div', class_='post-date')
                try: parsed_date = parser.parse(date_tag.get_text(strip=True))
                except: continue
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "CEF"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

## - Discursos - Banco de España - 
@st.cache_data(show_spinner=False)
def load_data_bde(start_date_str, end_date_str):
    """Extractor Banco de España - Versión con extracción de nombres reales desde PDF"""
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from PyPDF2 import PdfReader
    import io
    import requests
    import datetime
    import time
    import re

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BdE (España): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://www.bde.es/wbe/en/noticias-eventos/actualidad-banco-espana/intervenciones-publicas/"
    
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    
    def extraer_autor_y_cargo_desde_pdf(pdf_url):
        """Extrae el nombre y cargo del autor desde el PDF"""
        try:
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(pdf_url, headers=headers, timeout=15)
            if response.status_code != 200:
                return None, None
            
            pdf_file = io.BytesIO(response.content)
            reader = PdfReader(pdf_file)
            
            text = ""
            for i in range(min(3, len(reader.pages))):
                page_text = reader.pages[i].extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if not text:
                return None, None
            
            lineas = text.split('\n')
            nombre = None
            cargo = None
            
            for i, linea in enumerate(lineas):
                linea_limpia = linea.strip()
                
                if re.search(r'Governor|Gobernador', linea_limpia, re.IGNORECASE):
                    cargo = "Governor"
                    if i > 0 and lineas[i-1].strip() and len(lineas[i-1].strip().split()) >= 2:
                        nombre = lineas[i-1].strip()
                    elif i + 1 < len(lineas) and lineas[i+1].strip() and len(lineas[i+1].strip().split()) >= 2:
                        nombre = lineas[i+1].strip()
                    break
                elif re.search(r'Deputy Governor|Subgobernador', linea_limpia, re.IGNORECASE):
                    cargo = "Deputy Governor"
                    if i > 0 and lineas[i-1].strip() and len(lineas[i-1].strip().split()) >= 2:
                        nombre = lineas[i-1].strip()
                    elif i + 1 < len(lineas) and lineas[i+1].strip() and len(lineas[i+1].strip().split()) >= 2:
                        nombre = lineas[i+1].strip()
                    break
                elif re.search(r'Subgobernadora', linea_limpia, re.IGNORECASE):
                    cargo = "Subgobernadora"
                    if i > 0 and lineas[i-1].strip() and len(lineas[i-1].strip().split()) >= 2:
                        nombre = lineas[i-1].strip()
                    elif i + 1 < len(lineas) and lineas[i+1].strip() and len(lineas[i+1].strip().split()) >= 2:
                        nombre = lineas[i+1].strip()
                    break
                elif re.search(r'D\.G\.|Director General', linea_limpia, re.IGNORECASE):
                    cargo = "Director General"
                    if i > 0 and lineas[i-1].strip() and len(lineas[i-1].strip().split()) >= 2:
                        nombre = lineas[i-1].strip()
                    elif i + 1 < len(lineas) and lineas[i+1].strip() and len(lineas[i+1].strip().split()) >= 2:
                        nombre = lineas[i+1].strip()
                    break
            
            if not nombre:
                for linea in lineas[:15]:
                    linea_limpia = linea.strip()
                    if re.match(r'^[A-ZÁÉÍÓÚÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,}){1,3}$', linea_limpia):
                        if not any(palabra in linea_limpia for palabra in ['DIRECTOR', 'GENERAL', 'DEPARTAMENTO', 'SECRETARÍA', 'MINISTERIO', 'GOBIERNO', 'BANCO', 'ESPAÑA', 'MADRID']):
                            nombre = linea_limpia
                            break
            
            if nombre:
                nombre = ' '.join(nombre.split())
                nombre = nombre.title()
                nombre = re.sub(r'\bDe\b', 'de', nombre)
                nombre = re.sub(r'\bY\b', 'y', nombre)
                return nombre, cargo
            
            return None, None
            
        except Exception as e:
            print(f"      ⚠️ Error extrayendo del PDF: {e}")
            return None, None

    try:
        driver = webdriver.Chrome(options=chrome_options)
        driver.get(url)
        time.sleep(8)

        js_script = """
        let data = [];
        let results = document.querySelectorAll('.block-search-result, .block-search-result--image');
        results.forEach(el => {
            let titleEl = el.querySelector('.block-search-result__title, a');
            let dateEl = el.querySelector('.block-search-result__date');
            let linkEl = el.querySelector('a');
            if (titleEl && dateEl && linkEl) {
                data.push({
                    title: titleEl.innerText,
                    dateText: dateEl.innerText,
                    link: linkEl.href
                });
            }
        });
        return data;
        """
        extracted = driver.execute_script(js_script)
        driver.quit()

        print(f"   📚 Discursos encontrados: {len(extracted)}")

        for idx, item in enumerate(extracted):
            raw_title = item['title'].strip()
            raw_date_str = item['dateText'].strip()
            page_link = item['link']
            
            if not raw_title or not raw_date_str:
                continue

            parsed_date = None
            try:
                parsed_date = datetime.datetime.strptime(raw_date_str, '%d/%m/%Y')
            except:
                match = re.search(r'(\d{2}/\d{2}/\d{4})', raw_date_str)
                if match:
                    parsed_date = datetime.datetime.strptime(match.group(1), '%d/%m/%Y')

            if parsed_date and start_date <= parsed_date <= end_date:
                print(f"   🔍 Procesando ({idx+1}/{len(extracted)}): {parsed_date.strftime('%Y-%m-%d')}")
                
                try:
                    page_response = requests.get(page_link, headers={'User-Agent': 'Mozilla/5.0'}, timeout=10)
                    if page_response.status_code == 200:
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(page_response.text, 'html.parser')
                        pdf_link = None
                        for a in soup.find_all('a', href=True):
                            if a['href'].endswith('.pdf'):
                                pdf_link = a['href']
                                if pdf_link.startswith('/'):
                                    pdf_link = "https://www.bde.es" + pdf_link
                                break
                        
                        if pdf_link:
                            print(f"      📄 PDF encontrado, extrayendo autor...")
                            autor, cargo = extraer_autor_y_cargo_desde_pdf(pdf_link)
                            # Dentro de load_data_bde(), después de encontrar el autor
                            if autor:
                                titulo_limpio = raw_title
                                
                                # ========== NUEVA LIMPIEZA MEJORADA ==========
                                # Eliminar patrones comunes de cargo (en español e inglés)
                                patrones_cargo = [
                                    r'D\.G\.\s*Econom[íi]a\.\s*',      # D.G. Economía. o D.G. Economics.
                                    r'D\.G\.\s*Economics\.\s*',         # D.G. Economics.
                                    r'Deputy\s*Governor\.\s*',          # Deputy Governor.
                                    r'Governor\.\s*',                   # Governor.
                                    r'Subgobernador[a]?\.\s*',          # Subgobernadora. o Subgobernador.
                                    r'Director\s*General\.\s*',         # Director General.
                                    r'Head\s*of\s*\w+\.\s*',            # Head of Department.
                                    r'Director\.\s*',                   # Director.
                                    r'Chief\s*Economist\.\s*',          # Chief Economist.
                                    r'Gerente\s*General\.\s*',          # Gerente General.
                                    r'Vicepresident[ae]\.\s*',          # Vicepresidenta. o Vicepresidente.
                                    r'President[ae]\.\s*',              # Presidenta. o Presidente.
                                ]
                                
                                for patron in patrones_cargo:
                                    titulo_limpio = re.sub(patron, '', titulo_limpio, flags=re.IGNORECASE)
                                
                                # También eliminar cualquier texto entre paréntesis que parezca un cargo
                                titulo_limpio = re.sub(r'\s*\([^)]*(?:D\.G\.|Governor|Director|Econom[íi]a)[^)]*\)\s*', ' ', titulo_limpio, flags=re.IGNORECASE)
                                
                                # Limpiar espacios múltiples y puntos al inicio
                                titulo_limpio = re.sub(r'\s+', ' ', titulo_limpio).strip()
                                titulo_limpio = re.sub(r'^\.\s*', '', titulo_limpio)
                                
                                # Construir título final
                                titulo_final = f"{autor}: {titulo_limpio}"
                                
                                # Limpieza adicional: eliminar " : " si el título está vacío
                                titulo_final = re.sub(r':\s*$', '', titulo_final)
                                
                                print(f"      📝 Título limpio: {titulo_final[:80]}...")
                            
                            else:
                                print(f"      ⚠️ No se pudo extraer autor, usando formato estándar")
                                titulo_final = re.sub(r'\.\s+', ': ', raw_title, count=1)
                                titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                    else:
                        titulo_final = re.sub(r'\.\s+', ': ', raw_title, count=1)
                        
                except Exception as e:
                    print(f"      ⚠️ Error accediendo a la página: {e}")
                    titulo_final = re.sub(r'\.\s+', ': ', raw_title, count=1)
                
                if not any(r['Link'] == page_link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo_final,
                        "Link": page_link,
                        "Organismo": "BdE (España)"
                    })
                    print(f"      ✅ Agregado: {titulo_final[:80]}...")

    except Exception as e:
        print(f"❌ Error BDE: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    
    print(f"📊 BdE (España) - Total final: {len(df)}")
    return df

# ==========================================
# NUEVAS FUNCIONES PARA BID (bypass Cloudflare)
# ==========================================

@st.cache_data(show_spinner=False)
def load_investigacion_bid_cloudscraper(start_date_str, end_date_str):
    """
    Extrae Working Papers usando cloudscraper (bypass Cloudflare)
    """
    try:
        import cloudscraper
    except ImportError:
        print("❌ cloudscraper no instalado. Ejecuta: pip install cloudscraper")
        return pd.DataFrame()
    
    from bs4 import BeautifulSoup
    import datetime
    import re
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Cloudscraper: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    # Crear scraper con configuraciones específicas
    scraper = cloudscraper.create_scraper(
        browser={
            'browser': 'chrome',
            'platform': 'windows',
            'mobile': False
        },
        delay=5
    )
    
    # URLs a probar
    urls_to_try = [
        "https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers",
        "https://publications.iadb.org/es?f%5B0%5D=type%3A4633&f%5B1%5D=type%3ADocumentos%20de%20Trabajo"
    ]
    
    for url in urls_to_try:
        lang = "en" if "en?" in url else "es"
        try:
            print(f"📡 Accediendo a {url[:60]}...")
            response = scraper.get(url, timeout=30)
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Extraer artículos
                articles = soup.find_all('div', class_='views-row')
                print(f"   📚 Artículos encontrados: {len(articles)}")
                
                for article in articles:
                    # Extraer título y link
                    title_elem = article.find('div', class_='views-field-field-title')
                    if not title_elem:
                        continue
                    
                    a_tag = title_elem.find('a')
                    if not a_tag:
                        continue
                    
                    titulo = a_tag.get_text(strip=True)
                    link = a_tag.get('href')
                    if link and not link.startswith('http'):
                        link = "https://publications.iadb.org" + link
                    
                    # Extraer fecha
                    date_elem = article.find('div', class_='views-field-field-date-issued-text')
                    if date_elem:
                        date_text = date_elem.get_text(strip=True)
                        # Parsear fecha "Mar 2026"
                        match = re.search(r'([A-Za-z]{3})\s+(\d{4})', date_text)
                        if match:
                            mes_str, año = match.groups()
                            meses = {'Jan':1, 'Feb':2, 'Mar':3, 'Apr':4, 'May':5, 'Jun':6,
                                   'Jul':7, 'Aug':8, 'Sep':9, 'Oct':10, 'Nov':11, 'Dec':12}
                            mes = meses.get(mes_str, 1)
                            parsed_date = datetime.datetime(int(año), mes, 1)
                            
                            if start_date <= parsed_date <= end_date:
                                rows.append({
                                    "Date": parsed_date,
                                    "Title": titulo,
                                    "Link": link,
                                    "Organismo": f"BID ({'Inglés' if lang == 'en' else 'Español'})"
                                })
                                print(f"      ✅ {parsed_date.strftime('%Y-%m')}: {titulo[:50]}...")
            
        except Exception as e:
            print(f"⚠️ Error en {lang}: {e}")
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    
    print(f"📊 BID Cloudscraper - Total: {len(df)} documentos")
    return df


@st.cache_data(show_spinner=False)
def load_investigacion_bid_selenium_fallback(start_date_str, end_date_str):
    """
    Fallback: Extrae Working Papers con Selenium + delay largo
    """
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Selenium Fallback: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    chrome_options = Options()
    chrome_options.add_argument('--headless=new')
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')
    chrome_options.add_argument('--disable-blink-features=AutomationControlled')
    chrome_options.add_argument('--window-size=1920,1080')
    
    urls = [
        ("https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers", "en"),
        ("https://publications.iadb.org/es?f%5B0%5D=type%3A4633&f%5B1%5D=type%3ADocumentos%20de%20Trabajo", "es")
    ]
    
    for url, lang in urls:
        driver = None
        try:
            print(f"📡 Accediendo con Selenium a {url[:60]}...")
            driver = webdriver.Chrome(options=chrome_options)
            driver.get(url)
            
            # ⚠️ CLAVE: Esperar a que Cloudflare resuelva
            print("   ⏳ Esperando 20 segundos para Cloudflare...")
            time.sleep(20)
            
            # Scroll para cargar contenido
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(3)
            
            # Extraer usando BeautifulSoup
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            articles = soup.find_all('div', class_='views-row')
            print(f"   📚 Artículos encontrados: {len(articles)}")
            
            for article in articles:
                title_elem = article.find('div', class_='views-field-field-title')
                if not title_elem:
                    continue
                
                a_tag = title_elem.find('a')
                if not a_tag:
                    continue
                
                titulo = a_tag.get_text(strip=True)
                link = a_tag.get('href')
                if link and not link.startswith('http'):
                    link = "https://publications.iadb.org" + link
                
                date_elem = article.find('div', class_='views-field-field-date-issued-text')
                if date_elem:
                    date_text = date_elem.get_text(strip=True)
                    match = re.search(r'([A-Za-z]{3})\s+(\d{4})', date_text)
                    if match:
                        mes_str, año = match.groups()
                        meses = {'Jan':1, 'Feb':2, 'Mar':3, 'Apr':4, 'May':5, 'Jun':6,
                               'Jul':7, 'Aug':8, 'Sep':9, 'Oct':10, 'Nov':11, 'Dec':12,
                               'ene':1, 'feb':2, 'mar':3, 'abr':4, 'may':5, 'jun':6,
                               'jul':7, 'ago':8, 'sep':9, 'oct':10, 'nov':11, 'dic':12}
                        mes = meses.get(mes_str, 1)
                        parsed_date = datetime.datetime(int(año), mes, 1)
                        
                        if start_date <= parsed_date <= end_date:
                            rows.append({
                                "Date": parsed_date,
                                "Title": titulo,
                                "Link": link,
                                "Organismo": f"BID ({'Inglés' if lang == 'en' else 'Español'})"
                            })
            
        except Exception as e:
            print(f"⚠️ Error Selenium en {lang}: {e}")
        finally:
            if driver:
                driver.quit()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    
    print(f"📊 BID Selenium - Total: {len(df)} documentos")
    return df


def load_investigacion_bid_unified(start_date_str, end_date_str):
    """
    UNIFICADOR: Prueba cloudscraper primero, si falla usa Selenium
    """
    print("="*50)
    print("🔍 Iniciando extracción BID con estrategia unificada")
    print("="*50)
    
    # Intentar primero con cloudscraper
    try:
        print("\n🚀 Estrategia 1: Cloudscraper")
        df = load_investigacion_bid_cloudscraper(start_date_str, end_date_str)
        if not df.empty:
            print(f"✅ Cloudscraper exitoso: {len(df)} documentos")
            return df
        else:
            print("⚠️ Cloudscraper no obtuvo resultados")
    except Exception as e:
        print(f"⚠️ Cloudscraper falló: {e}")
    
    # Fallback a Selenium
    print("\n🚀 Estrategia 2: Selenium con delay largo")
    try:
        df = load_investigacion_bid_selenium_fallback(start_date_str, end_date_str)
        if not df.empty:
            print(f"✅ Selenium exitoso: {len(df)} documentos")
            return df
        else:
            print("⚠️ Selenium no obtuvo resultados")
    except Exception as e:
        print(f"⚠️ Selenium falló: {e}")
    
    print("\n❌ Ambas estrategias fallaron para BID")
    return pd.DataFrame()


# ==========================================
# EXPORTACIÓN A WORD
# ==========================================
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    c = docx.oxml.shared.OxmlElement('w:color'); c.set(docx.oxml.shared.qn('w:val'), '0000EE'); rPr.append(c)
    u = docx.oxml.shared.OxmlElement('w:u'); u.set(docx.oxml.shared.qn('w:val'), 'single'); rPr.append(u)
    b = docx.oxml.shared.OxmlElement('w:b'); rPr.append(b)
    
    for s in ['w:sz', 'w:szCs']:
        sz = docx.oxml.shared.OxmlElement(s); sz.set(docx.oxml.shared.qn('w:val'), '28'); rPr.append(sz)
        
    rFonts = docx.oxml.shared.OxmlElement('w:rFonts'); rFonts.set(docx.oxml.shared.qn('w:ascii'), 'Calibri'); rFonts.set(docx.oxml.shared.qn('w:hAnsi'), 'Calibri'); rPr.append(rFonts)
    t = docx.oxml.shared.OxmlElement('w:t'); t.text = text; new_run.append(rPr); new_run.append(t); hyperlink.append(new_run); paragraph._p.append(hyperlink)

def generate_word(df, title="Boletín Mensual", subtitle=""):
    doc = Document()
    h = doc.add_heading(title, 0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle); run.font.name, run.font.size = 'Calibri', Pt(14)
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=len(df.columns)-1)
    table.style = 'Table Grid'
    
    cols = [c for c in df.columns if c != 'Link']
    
    for idx, name in enumerate(cols):
        p = table.rows[0].cells[idx].paragraphs[0]
        run = p.add_run(name)
        run.font.name = 'Calibri'
        run.font.size = Pt(14) 
        run.bold = True
        
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            p = cells[i].paragraphs[0]
            if col == 'Nombre de Documento': 
                add_hyperlink(p, str(row[col]), str(row['Link']))
            else:
                run = p.add_run(str(row[col]))
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.bold = True

    if 'Tipo de Documento' in df.columns and 'Organismo' in df.columns:
        col_tipo = cols.index('Tipo de Documento')
        col_org = cols.index('Organismo')
        
        start_row = 1
        while start_row <= len(df):
            cat_val = df.iloc[start_row - 1]['Tipo de Documento']
            org_val = df.iloc[start_row - 1]['Organismo']
            end_row = start_row
            
            if cat_val == "Discursos":
                table.cell(start_row, col_org).text = "" 
                while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == "Discursos":
                    table.cell(end_row + 1, col_org).text = "" 
                    end_row += 1
                
                if end_row > start_row:
                    target_cell = table.cell(start_row, col_org)
                    target_cell.merge(table.cell(end_row, col_org))
                
                start_row = end_row + 1
                continue
                
            while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == cat_val and df.iloc[end_row]['Organismo'] == org_val:
                table.cell(end_row + 1, col_org).text = "" 
                end_row += 1
                
            if end_row > start_row:
                target_cell = table.cell(start_row, col_org)
                target_cell.merge(table.cell(end_row, col_org))
                target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                
            start_row = end_row + 1

        start_row = 1
        while start_row <= len(df):
            cat_val = df.iloc[start_row - 1]['Tipo de Documento']
            end_row = start_row
            
            while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == cat_val:
                table.cell(end_row + 1, col_tipo).text = ""
                end_row += 1
            
            if end_row > start_row:
                target_cell = table.cell(start_row, col_tipo)
                target_cell.merge(table.cell(end_row, col_tipo))
                target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                
            start_row = end_row + 1
                
    out = BytesIO(); doc.save(out); out.seek(0); return out

# ==========================================
# INTERFAZ DE USUARIO Y MAIN
# ==========================================
try: 
    st.sidebar.image("logo_banxico.png", use_container_width=True)
except: 
    st.sidebar.markdown("### 🏦 BANCO DE MÉXICO")

st.sidebar.markdown("---")
st.sidebar.header("Menú de Navegación")
modo_app = st.sidebar.radio("", ["Boletín", "Categorías"], key="menu_principal") 
st.sidebar.markdown("---")

anios_str = ["2026", "2025", "2024", "2023", "2022"]
meses_dict = {
    "Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6,
    "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12
}

# --- LISTAS DINÁMICAS DE ORGANISMOS ---
orgs_discursos = ["BBk (Alemania)", "BdE (España)", "BdF (Francia)", "BM", "BoC (Canadá)", "BoE (Inglaterra)", "BoJ (Japón)", "BPI", "CEF", "ECB (Europa)", "Fed (Estados Unidos)", "FMI", "PBoC (China)"]
orgs_reportes = ["BID", "BM", "BPI", "CEF", "FEM", "OCDE"]
orgs_pub_inst = ["BM", "BPI", "CEF", "CEMLA", "FMI", "FMI (Mission Concluding)", "F&D", "G20", "OCDE", "OEI", "F&D Magazine"] 
orgs_investigacion = ["BID", "BM", "BPI", "CEMLA", "FMI", "OCDE"]

if modo_app == "Boletín":
    st.title("Generador de Boletín Mensual")
    st.markdown("Extrae y unifica documentos de todas las categorías y organismos por mes."); st.markdown("---")
    
    c1, c2 = st.columns(2)
    m_sel = c1.multiselect("Mes(es)", options=list(meses_dict.keys()))
    a_sel = c2.multiselect("Año(s)", options=anios_str, default=["2026"])
    
    if st.button("📄 Generar Boletín Mensual", type="primary"):
        if not m_sel or not a_sel: 
            st.warning("⚠️ Selecciona mes y año.")
        else:
            m_num = [meses_dict[m] for m in m_sel]
            a_num = [int(a) for a in a_sel]
            sd = f"01.{min(m_num):02d}.{min(a_num)}"
            ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"
            
            all_dfs = []
            prog = st.progress(0)
            txt = st.empty()
            
            total_pasos = len(orgs_discursos) + len(orgs_reportes) + len(orgs_pub_inst) + len(orgs_investigacion)
            paso_actual = 0
            
            # 1. BARRIDO DE DISCURSOS
            for org in orgs_discursos:
                txt.text(f"Procesando Discursos: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BPI": df = load_data_bis()
                    elif org == "ECB (Europa)": df = load_data_ecb(sd, ed)
                    elif org == "BBk (Alemania)": df = load_data_bbk(sd, ed)
                    elif org == "Fed (Estados Unidos)": df = load_data_fed(a_num)
                    elif org == "BdF (Francia)": df = load_data_bdf(sd, ed)
                    elif org == "BM": df = load_data_bm(sd, ed)
                    elif org == "BoC (Canadá)": df = load_data_boc(sd, ed)
                    elif org == "BoJ (Japón)": df = load_data_boj(sd, ed)
                    elif org == "CEF": df = load_data_cef(sd, ed)
                    elif org == "PBoC (China)": df = load_data_pboc(sd, ed)
                except Exception as e: pass 
                
                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty: 
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Discursos"
                        all_dfs.append(df_f)
                paso_actual += 1; prog.progress(paso_actual / total_pasos)

            # 2. BARRIDO DE REPORTES
            for org in orgs_reportes:
                txt.text(f"Procesando Reportes: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BID": df = load_reportes_bid_en(sd, ed)
                    elif org == "BM": df = load_reportes_bm(sd, ed) # <--- AGRÉGALO AQUÍ
                    elif org == "BPI": df = load_reportes_bpi(sd, ed)
                    elif org == "CEF": df = load_reportes_cef(sd, ed)
                    elif org == "OCDE": df = load_reportes_ocde(sd, ed)
                except Exception as e: pass
                
                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty: 
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Reportes"
                        all_dfs.append(df_f)
                paso_actual += 1; prog.progress(paso_actual / total_pasos)
                
            # 3. BARRIDO DE PUBLICACIONES INSTITUCIONALES 
            for org in orgs_pub_inst:
                txt.text(f"Procesando Pub. Institucionales: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BPI": df = load_pub_inst_bpi(sd, ed)
                    elif org == "CEF": df = load_pub_inst_cef(sd, ed)
                    elif org == "BM": df = load_pub_inst_bm(sd, ed)
                    elif org == "FMI": 
                        # 1. SSG - JSON Estático (WEO, Fiscal Monitor)
                        df_flagships = load_pub_inst_fmi(sd, ed)
                        
                        # 2. SSG - JSON Estático (Comunicados)
                        df_prs = load_press_releases_fmi(sd, ed)
                        
                        # 3. CSR API - Coveo (Country Reports)
                        df_crs = load_country_reports_fmi(sd, ed) # <-- LA NUEVA API
                        
                        # Unión
                        dfs_a_unir = [d for d in [df_flagships, df_prs, df_crs] if not d.empty]
                        if dfs_a_unir:
                            df = pd.concat(dfs_a_unir, ignore_index=True)
                            df = df.sort_values("Date", ascending=False)
                            
                except Exception as e: pass 
                
                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty: 
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Publicaciones Institucionales"
                        all_dfs.append(df_f)
                paso_actual += 1; prog.progress(paso_actual / total_pasos)

            # 4. BARRIDO DE INVESTIGACIÓN
            for org in orgs_investigacion:
                txt.text(f"Procesando Investigación: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BPI": df = load_investigacion_bpi(sd, ed)
                    elif org == "BM": df = load_investigacion_bm(sd, ed)
                except Exception as e: pass 
                
                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty: 
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Investigación"
                        all_dfs.append(df_f)
                paso_actual += 1; prog.progress(paso_actual / total_pasos)
            
            txt.empty()
            prog.empty()
            
            # --- CONSOLIDACIÓN FINAL ---
            if all_dfs:
                f_df = pd.concat(all_dfs, ignore_index=True)
                
                # 1. SEPARAR Y ORDENAR CON REGLAS Y JERARQUÍA ESTRICTA
                df_rep = f_df[f_df['Categoría'] == "Reportes"].copy()
                df_pub = f_df[f_df['Categoría'] == "Publicaciones Institucionales"].copy()
                df_inv = f_df[f_df['Categoría'] == "Investigación"].copy()
                df_disc = f_df[f_df['Categoría'] == "Discursos"].copy()
                
                # Ordenamiento específico
                if not df_rep.empty: df_rep = df_rep.sort_values(by=["Organismo", "Title"], ascending=[True, True])
                if not df_pub.empty: df_pub = df_pub.sort_values(by=["Organismo", "Title"], ascending=[True, True])
                if not df_inv.empty: df_inv = df_inv.sort_values(by=["Organismo", "Title"], ascending=[True, True])
                if not df_disc.empty: df_disc = df_disc.sort_values(by=["Title"], ascending=[True]) # Sin agrupar por organismo
                
                # Unimos respetando tu jerarquía exacta
                f_df = pd.concat([df_rep, df_pub, df_inv, df_disc], ignore_index=True)
                
                # 2. COLUMNAS: Dejamos las 3 solicitadas + Link
                f_df = f_df[['Categoría', 'Organismo', 'Title', 'Link']]
                f_df = f_df.rename(columns={"Categoría": "Tipo de Documento", "Title": "Nombre de Documento"})
                
                st.success(f"Se consolidaron **{len(f_df)}** documentos en total.")
                word = generate_word(f_df, subtitle=", ".join(m_sel) + " " + ", ".join(a_sel))
                st.download_button("📄 Descargar Boletín", word, f"Boletin_{'_'.join(m_sel)}.docx")
                
                disp = f_df.copy()
                disp["Nombre de Documento"] = disp.apply(lambda x: f"[{x['Nombre de Documento']}]({x['Link']})", axis=1)
                st.markdown(disp[["Tipo de Documento", "Organismo", "Nombre de Documento"]].to_markdown(index=False), unsafe_allow_html=True)
            else: 
                st.warning("No se encontraron documentos para los criterios seleccionados.")

elif modo_app == "Categorías":
    st.title("Documentos de Organismos Internacionales")
    tipo_doc = st.sidebar.selectbox("Tipo de Documento", ["Discursos", "Reportes", "Investigación", "Publicaciones Institucionales"])
    
    # Construcción segura de las listas de interfaz
    if tipo_doc == "Discursos": orgs_list = ["Todos"] + sorted(orgs_discursos)
    elif tipo_doc == "Reportes": orgs_list = ["Todos"] + sorted(orgs_reportes)
    elif tipo_doc == "Investigación": orgs_list = ["Todos"] + sorted(orgs_investigacion)
    elif tipo_doc == "Publicaciones Institucionales": orgs_list = ["Todos"] + sorted(orgs_pub_inst)
    else: orgs_list = ["Todos"] + sorted(list(set(orgs_discursos + orgs_reportes + orgs_investigacion + orgs_pub_inst)))
        
    organismo_seleccionado = st.sidebar.selectbox("Organismo", orgs_list)
    
    c1, c2 = st.columns(2)
    m_sel = c1.multiselect("Mes(es)", options=list(meses_dict.keys()))
    a_sel = c2.multiselect("Año(s)", options=anios_str, default=["2026"])
    
    if st.button("🔍 Buscar", type="primary"):
        if not m_sel or not a_sel:
            st.warning("⚠️ Selecciona mes y año.")
        else:
            m_num = [meses_dict[m] for m in m_sel]
            a_num = [int(a) for a in a_sel]
            sd = f"01.{min(m_num):02d}.{min(a_num)}"
            ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"
            
            target_orgs = orgs_list[1:] if organismo_seleccionado == "Todos" else [organismo_seleccionado]
            dfs_comb = []
            progreso = st.progress(0)
            txt = st.empty()
            
            for i, o in enumerate(target_orgs):
                txt.text(f"Extrayendo: {o}...")
                df = pd.DataFrame()
                try:
                    # --- LÓGICA DE EXTRACCIÓN POR TIPO ---
                    if tipo_doc == "Discursos":
                        if o == "BPI": df = load_data_bis()
                        elif o == "ECB (Europa)": df = load_data_ecb(sd, ed)
                        elif o == "BBk (Alemania)": df = load_data_bbk(sd, ed)
                        elif o == "Fed (Estados Unidos)": df = load_data_fed(a_num)
                        elif o == "BdF (Francia)": df = load_data_bdf(sd, ed)
                        elif o == "BM": df = load_data_bm(sd, ed)
                        elif o == "BoC (Canadá)": df = load_data_boc(sd, ed)
                        elif o == "BoJ (Japón)": df = load_data_boj(sd, ed)
                        elif o == "BoE (Inglaterra)": df = load_discursos_boe(sd, ed)
                        elif o == "CEMLA": 
                            print("🔴🔴🔴 LLAMANDO A CEMLA INVESTIGACIÓN 🔴🔴🔴")
                            df = load_investigacion_cemla(sd, ed)
                            print(f"🔴🔴🔴 RESULTADO CEMLA: {len(df)} documentos 🔴🔴🔴")
                        elif o == "CEF": df = load_data_cef(sd, ed)
                        elif o == "FMI":
                            df = load_discursos_fmi(sd, ed)
                        elif o == "PBoC (China)": df = load_data_pboc(sd, ed)
                        elif o == "BdE (España)":
                            df = load_data_bde(sd, ed)
                    
                    elif tipo_doc == "Reportes":
                        if o == "BID": 
                            dfs_bid = []
                            try:
                                dfs_bid.append(load_reportes_bid(sd, ed))
                            except:
                                pass
                            try:
                                dfs_bid.append(load_reportes_bid_en(sd, ed))
                            except:
                                pass
                            dfs_bid = [d for d in dfs_bid if not d.empty]
                            if dfs_bid:
                                df = pd.concat(dfs_bid, ignore_index=True).drop_duplicates(
                                    subset=['Link'])
                        elif o == "BM": df = load_reportes_bm(sd, ed) # <--- AGRÉGALO AQUÍ
                        elif o == "BPI": df = load_reportes_bpi(sd, ed)
                        elif o == "CEF": df = load_reportes_cef(sd, ed)
                        elif o == "OCDE": df = load_reportes_ocde(sd, ed)
                        elif o == "FEM": df = load_reportes_fem(sd, ed)
                        
                    elif tipo_doc == "Investigación":
                        if o == "BID":
                            df = load_investigacion_bid_unified(sd, ed)
                        elif o == "BPI": df = load_investigacion_bpi(sd, ed)
                        elif o == "BM": df = load_investigacion_bm(sd, ed)
                        elif o == "CEMLA":   # <-- ESTA LÍNEA DEBE EXISTIR
                            print("🔴 LLAMANDO A CEMLA")
                            df = load_investigacion_cemla(sd, ed)
                        elif o == "FMI": 
                            df_blogs, df_wp = pd.DataFrame(), pd.DataFrame()
                            try: df_blogs = load_investigacion_fmi(sd, ed)
                            except: pass
                            try: df_wp = load_working_papers_fmi(sd, ed)
                            except: pass
                            
                            dfs_fmi = [d for d in [df_blogs, df_wp] if not d.empty]
                            if dfs_fmi:
                                df = pd.concat(dfs_fmi, ignore_index=True).drop_duplicates(subset=['Link']).sort_values("Date", ascending=False)
                        elif o == "OCDE":
                            df = load_investigacion_ocde(sd, ed)
                        
                    elif tipo_doc == "Publicaciones Institucionales":
                        if o == "BPI": 
                            df = load_pub_inst_bpi(sd, ed)
                        elif o == "CEF": 
                            df = load_pub_inst_cef(sd, ed)
                        elif o == "OEI": 
                            df = load_pub_inst_oei(sd, ed)
                        elif o == "OCDE":
                            df = load_pub_inst_ocde(sd, ed)
                        elif o == "BM": 
                            df = load_pub_inst_bm(sd, ed)
                        elif o == "CEMLA":
                            df = load_pub_inst_cemla(sd, ed)
                        elif o == "FMI":
                            print(f"\n{'='*50}")
                            print(f"🔍 CATEGORÍAS - Procesando FMI para {m_sel} {a_sel}")
                            print(f"   Fechas: {sd} a {ed}")
                            print(f"{'='*50}")
                            
                            # 1. F&D Magazine (agregar explícitamente)
                            df_fandd = load_pub_inst_fandd(sd, ed)
                            print(f"   📊 F&D Magazine: {len(df_fandd)} documentos")

                            # 2. SSG - JSON Estático (WEO, Fiscal Monitor)
                            df_flagships = load_pub_inst_fmi(sd, ed)
                            print(f"   📊 Flagships: {len(df_flagships)} documentos")
                            
                            # 3. SSG - JSON Estático (Comunicados)
                            df_prs = load_press_releases_fmi(sd, ed)
                            print(f"   📊 Press Releases: {len(df_prs)} documentos")
                            
                            # 4. CSR API - Coveo (Country Reports)
                            df_crs = load_country_reports_fmi(sd, ed)
                            print(f"   📊 Country Reports: {len(df_crs)} documentos")
                            
                            # Unir todos
                            print(f"🔍 CATEGORÍAS - Flagships: {len(df_flagships)}, PRs: {len(df_prs)}, CRs: {len(df_crs)}")
                            dfs_a_unir = [d for d in [df_fandd, df_flagships, df_prs, df_crs] if not d.empty]
                            if dfs_a_unir:
                                df = pd.concat(dfs_a_unir, ignore_index=True)
                                df = df.sort_values("Date", ascending=False)
                                print(f"   📊 TOTAL combinado FMI: {len(df)} documentos")
                            else:
                                print(f"   ⚠️ Ninguna fuente retornó datos")
                        
                except Exception as e:
                    pass
                
                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty: 
                        df_f['Organismo'] = o
                        dfs_comb.append(df_f)
                progreso.progress((i+1)/len(target_orgs))
            
            txt.empty()
            progreso.empty()
            
            if dfs_comb:
                f_df = pd.concat(dfs_comb, ignore_index=True)              
                # --- FORMATO HOMOGÉNEO (IGUAL AL BOLETÍN) ---
                f_df['Categoría'] = tipo_doc

                # ========== ELIMINACIÓN DE DUPLICADOS ==========
                print(f"📊 Total antes de desduplicar: {len(f_df)}")

                # 1. Eliminar duplicados exactos por Link
                f_df = f_df.drop_duplicates(subset=['Link'], keep='first')
                print(f"   Después de eliminar duplicados por Link: {len(f_df)}")

                # 2. Normalizar títulos para comparación
                f_df['Title_Norm'] = f_df['Title'].str.lower().str.replace(r'[^\w\s]', '', regex=True).str.replace(r'\s+', ' ', regex=True).str.strip()

                # 3. Eliminar duplicados por título normalizado (mismo título en diferentes categorías)
                f_df = f_df.sort_values('Date', ascending=False).drop_duplicates(subset=['Title_Norm'], keep='first')
                print(f"   Después de eliminar duplicados por título: {len(f_df)}")

                # 4. Eliminar columna temporal
                f_df = f_df.drop(columns=['Title_Norm'])

                print(f"📊 Total después de desduplicación: {len(f_df)}")

                # --- PREPARACIÓN PARA LA VISTA PREVIA ---
                disp = f_df.copy()
                disp = disp.sort_values(
                    by="Date", ascending=False)  # Orden cronológico
                disp["Fecha"] = disp["Date"].dt.strftime(
                    '%d/%m/%Y')  # Formatear fecha
                disp["Nombre de Documento"] = disp.apply(
                    lambda x: f"[{x['Title']}]({x['Link']})", axis=1)
                disp = disp.rename(columns={"Categoría": "Tipo de Documento"})

                if organismo_seleccionado == "Todos":
                    cols_vis = ["Fecha", "Tipo de Documento",
                                "Organismo", "Nombre de Documento"]
                else:
                    cols_vis = ["Fecha", "Tipo de Documento",
                                "Nombre de Documento"]

                st.markdown(disp[cols_vis].to_markdown(index=False), unsafe_allow_html=True)
            else:
                # Verificar si CEMLA estaba en la búsqueda y no hay resultados
                if "CEMLA" in target_orgs and tipo_doc == "Investigación":
                    st.warning("⚠️ No se encontraron documentos para las fechas seleccionadas.")
                    st.info("📌 **CEMLA Investigación**: ScienceDirect bloquea el acceso automatizado. No se pueden extraer artículos.\n\n➡️ **Solución**: Utiliza la sección **'Carga Manual'** en el menú principal para agregar estos documentos al boletín mensual.")
                else:
                    st.warning(
                    "No se encontraron documentos para las fechas seleccionadas.")

elif modo_app == "Carga Manual":
    st.title("🛠️ Centro de Carga Manual")
    st.markdown("Pega el texto de las páginas que fallan. Previsualiza, valida y une todo en un solo documento.")
    
    if 'cargas_validadas' not in st.session_state:
        st.session_state.cargas_validadas = {
            "OCDE (Reportes)": pd.DataFrame(),
            "OCDE (Pub. Institucionales)": pd.DataFrame(),
            "OCDE (Investigación)": pd.DataFrame()
        }

    st.subheader("Estado de Carga")
    cols_estado = st.columns(3)
    claves_cajas = list(st.session_state.cargas_validadas.keys())
    
    for i, clave in enumerate(claves_cajas):
        estado = "✅ Listo" if not st.session_state.cargas_validadas[clave].empty else "❌ Pendiente"
        cols_estado[i].info(f"**{clave}**\n\n{estado}")

    st.markdown("---")
    
    c1, c2 = st.columns(2)
    mes_manual = c1.selectbox("Mes objetivo a filtrar:", [1,2,3,4,5,6,7,8,9,10,11,12], index=datetime.datetime.now().month-1, format_func=lambda x: calendar.month_name[x].capitalize())
    año_manual = c2.number_input("Año objetivo a filtrar:", min_value=2020, max_value=2030, value=datetime.datetime.now().year)

    st.markdown("---")
    st.subheader("Cajas de Extracción")

    def crear_caja_manual(titulo_caja, categoria_doc, organismo_nombre, url_fuente=None):
        with st.expander(f"📥 Cargar: {titulo_caja}", expanded=True):
            
            if url_fuente:
                st.markdown(f"👉 **[Haz clic aquí para abrir la página oficial de {titulo_caja}]({url_fuente})**")
                
            texto = st.text_area(f"Copia el texto de la página y pégalo aquí (Ctrl+A, Ctrl+C, Ctrl+V):", height=150, key=f"txt_{titulo_caja}")
            
            col_btn1, col_btn2 = st.columns([1, 1])
            
            if col_btn1.button(f"🔍 Previsualizar {titulo_caja}", key=f"btn_prev_{titulo_caja}"):
                if texto:
                    with st.spinner("Procesando y buscando links..."):
                        df_bruto = procesar_texto_pegado(texto, organismo_nombre)
                            
                        if not df_bruto.empty:
                            df_filtrado = df_bruto[
                                (df_bruto['Date'].dt.month == mes_manual) & 
                                (df_bruto['Date'].dt.year == año_manual)
                            ].copy()
                            
                            if not df_filtrado.empty:
                                for idx in df_filtrado.index:
                                    t = df_filtrado.loc[idx, "Title"]
                                    df_filtrado.loc[idx, "Link"] = buscar_link_inteligente(t, organismo_nombre)
                                
                                df_filtrado['Categoría'] = categoria_doc
                                st.session_state[f"temp_{titulo_caja}"] = df_filtrado
                                
                                st.success(f"Se encontraron {len(df_filtrado)} documentos de {mes_manual} {año_manual}.")
                                st.dataframe(df_filtrado, use_container_width=True)
                            else:
                                st.warning("No hay coincidencias con el mes y año seleccionados.")
                else:
                    st.error("Pega el texto primero.")
            
            if col_btn2.button(f"➕ Agregar a Descarga Final", type="primary", key=f"btn_add_{titulo_caja}"):
                if f"temp_{titulo_caja}" in st.session_state and not st.session_state[f"temp_{titulo_caja}"].empty:
                    st.session_state.cargas_validadas[titulo_caja] = st.session_state[f"temp_{titulo_caja}"]
                    st.success(f"¡{titulo_caja} guardado en memoria! ✅")
                    time.sleep(1)
                    st.rerun() 
                else:
                    st.error("Primero debes Previsualizar y obtener resultados.")

    link_ocde_rep = "https://www.oecd.org/en/search/publications.html?orderBy=mostRecent&page=0&facetTags=oecd-content-types%3Apublications%2Freports%2Coecd-languages%3Aen&minPublicationYear=2026&maxPublicationYear=2026"
    link_ocde_pub = "https://www.oecd.org/en/search.html?orderBy=mostRecent&page=0&facetTags=oecd-policy-subissues%3Apsi114%2Coecd-languages%3Aen"
    link_ocde_inv = "https://www.oecd.org/en/publications/reports.html?orderBy=mostRecent&page=0&facetTags=oecd-content-types%3Apublications%2Fworking-papers%2Coecd-languages%3Aen"
    
    crear_caja_manual("OCDE (Reportes)", "Reportes", "OCDE", link_ocde_rep)
    crear_caja_manual("OCDE (Pub. Institucionales)", "Publicaciones Institucionales", "OCDE", link_ocde_pub)
    crear_caja_manual("OCDE (Investigación)", "Investigación", "OCDE", link_ocde_inv)

    st.markdown("---")
    st.subheader("Exportación Final")
    
    tablas_listas = [df for df in st.session_state.cargas_validadas.values() if not df.empty]
    
    if tablas_listas:
        df_maestro = pd.concat(tablas_listas, ignore_index=True)
        num_cat = len(tablas_listas)
        st.info(f"Tienes **{num_cat}/3** categorías listas, sumando un total de **{len(df_maestro)}** documentos para exportar.")
        
        df_word_manual = df_maestro[['Categoría', 'Organismo', 'Title', 'Link']].copy()
        df_word_manual = df_word_manual.rename(columns={"Categoría": "Tipo de Documento", "Title": "Nombre de Documento"})
        
        word_manual = generate_word(df_word_manual, title="Boletín - Carga Manual", subtitle=f"Mes: {mes_manual} | Año: {año_manual}")
        
        c_down, c_clear = st.columns(2)
        with c_down:
            st.download_button(
                label=f"📄 Descargar Word ({num_cat}/3 Listas)", 
                data=word_manual, 
                file_name=f"Carga_Manual_{mes_manual}_{año_manual}.docx"
            )
        with c_clear:
            if st.button("🗑️ Reiniciar todo el módulo"):
                for clave in st.session_state.cargas_validadas.keys():
                    st.session_state.cargas_validadas[clave] = pd.DataFrame()
                st.rerun()
    else:
        st.warning("Aún no has agregado ninguna carga a la descarga final. Agrega al menos 1 para habilitar el botón de descarga.")

# ==========================================
# CÓDIGO DE PRUEBA (agregar al final de app.py)
# ==========================================
